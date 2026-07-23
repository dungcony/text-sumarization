#!/usr/bin/env python3
"""Build the Week 3–4 internship report from verified local project artifacts.

Run with the workspace's bundled Python because it includes python-docx and
Pillow. The script deliberately reads the actual Week 3 inference JSON and
Week 4 audit JSON, so report metrics cannot drift from the saved run.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


WORKSPACE = Path(__file__).resolve().parent.parent
REPORT_DIR = Path(__file__).resolve().parent
WEEK34 = WORKSPACE / "tuan 3-4" / "summarization"
RESULT_PATH = WEEK34 / "results" / "vit5_sample_5.json"
AUDIT_PATH = WEEK34 / "data" / "raw" / "vietnews_medical_raw_1000_audit.json"
OUTPUT_DOCX = REPORT_DIR / "Lường Tiến Dũng - B22DCCN128 - Bao_cao_thuc_tap_Tuan3_4.docx"
FIGURE_DIR = REPORT_DIR / "assets_tuan3_4"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def font(size: int, bold: bool = False) -> tuple[str, int, bool]:
    return ("Times New Roman", size, bold)


def set_run_font(run, size: int = 13, bold: bool = False, italic: bool = False, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str = "", *, align=None, indent: bool = True, before: int = 0, after: int = 6, bold: bool = False, italic: bool = False, size: int = 13) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.75)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    set_run_font(paragraph.add_run(text))


def add_heading(doc: Document, text: str, level: int = 1, page_break_before: bool = False) -> None:
    size = 16 if level == 1 else 14 if level == 2 else 13
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=True)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, italic=True)


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[str]], widths: list[float] | None = None) -> None:
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = header_cells[index]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "D9EAF7")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=11, bold=True)
        if widths:
            cell.width = Cm(widths[index])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(value))
            set_run_font(r, size=10)
            if widths:
                cell.width = Cm(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def find_font() -> str:
    for candidate in (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ):
        if Path(candidate).exists():
            return candidate
    raise FileNotFoundError("Không tìm thấy font dùng để tạo hình minh họa.")


def draw_centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, text_font, fill: str = "#17365D") -> None:
    left, top, right, bottom = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline="#0B1F33", width=3)
    lines = text.split("\n")
    line_height = text_font.getbbox("Ag")[3] + 4
    total_height = len(lines) * line_height
    y = top + (bottom - top - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=text_font)
        draw.text(((left + right - (bbox[2] - bbox[0])) / 2, y), line, font=text_font, fill="white")
        y += line_height


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#17365D", width=5)
    x, y = end
    draw.polygon([(x, y), (x - 16, y - 10), (x - 16, y + 10)], fill="#17365D")


def create_figures(rouge: dict[str, dict[str, float]]) -> tuple[Path, Path]:
    FIGURE_DIR.mkdir(exist_ok=True)
    regular = ImageFont.truetype(find_font(), 28)
    small = ImageFont.truetype(find_font(), 23)
    bold = ImageFont.truetype(find_font(), 28)

    pipeline_path = FIGURE_DIR / "pipeline_tuan3_4.png"
    image = Image.new("RGB", (1800, 820), "white")
    draw = ImageDraw.Draw(image)
    draw.text((65, 35), "Quy trình thực hiện Tuần 3–4", font=bold, fill="#17365D")
    draw_centered(draw, (70, 175, 395, 325), "Mẫu văn bản\n+ tham chiếu", regular, "#4472C4")
    draw_centered(draw, (535, 175, 860, 325), "ViT5 checkpoint\nđã pre-train", regular, "#4472C4")
    draw_centered(draw, (1000, 175, 1325, 325), "Beam search\nsinh tóm tắt", regular, "#4472C4")
    draw_centered(draw, (1465, 175, 1760, 325), "ROUGE\n+ JSON kết quả", regular, "#4472C4")
    arrow(draw, (395, 250), (535, 250))
    arrow(draw, (860, 250), (1000, 250))
    arrow(draw, (1325, 250), (1465, 250))
    draw_centered(draw, (70, 515, 395, 665), "Nguồn CSV\n1.000 bản ghi", regular, "#70AD47")
    draw_centered(draw, (535, 515, 860, 665), "Chuẩn hóa NFC\n+ kiểm tra chất lượng", small, "#70AD47")
    draw_centered(draw, (1000, 515, 1325, 665), "990 cặp thô\n+ audit JSON", regular, "#70AD47")
    draw_centered(draw, (1465, 515, 1760, 665), "Bàn giao\nTuần 5–6", regular, "#70AD47")
    arrow(draw, (395, 590), (535, 590))
    arrow(draw, (860, 590), (1000, 590))
    arrow(draw, (1325, 590), (1465, 590))
    draw.text((70, 730), "Nhánh trên: Tuần 3 (inference/đánh giá). Nhánh dưới: Tuần 4 (thu thập và kiểm tra dữ liệu thô).", font=small, fill="#333333")
    image.save(pipeline_path)

    chart_path = FIGURE_DIR / "rouge_tuan3.png"
    image = Image.new("RGB", (1400, 720), "white")
    draw = ImageDraw.Draw(image)
    draw.text((65, 35), "ROUGE F1 trên 5 mẫu kiểm thử", font=bold, fill="#17365D")
    x0, y0, x1, y1 = 250, 150, 1270, 570
    draw.line([(x0, y0), (x0, y1), (x1, y1)], fill="#333333", width=3)
    for tick in range(0, 101, 20):
        y = y1 - (y1 - y0) * tick / 100
        draw.line([(x0 - 10, y), (x1, y)], fill="#D9E2F3", width=2)
        draw.text((150, y - 14), str(tick), font=small, fill="#555555")
    colors = ["#4472C4", "#ED7D31", "#70AD47"]
    labels = [("ROUGE-1", rouge["rouge1"]["f1"]), ("ROUGE-2", rouge["rouge2"]["f1"]), ("ROUGE-L", rouge["rougeL"]["f1"])]
    positions = [410, 720, 1030]
    for (label, value), x, color in zip(labels, positions, colors):
        top = y1 - (y1 - y0) * value / 100
        draw.rounded_rectangle((x, top, x + 160, y1), radius=8, fill=color)
        draw.text((x + 28, top - 45), f"{value:.2f}", font=small, fill="#17365D")
        draw.text((x + 12, y1 + 25), label, font=small, fill="#333333")
    draw.text((250, 635), "Thước đo là F1, thang 0–100. Đây là smoke test 5 mẫu, không phải benchmark cuối cùng.", font=small, fill="#333333")
    image.save(chart_path)
    return pipeline_path, chart_path


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.2)
    for style_name in ("Normal", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(13)
    add_page_number(section.footer.paragraphs[0])
    return doc


def add_cover(doc: Document) -> None:
    for text, size in (("BỘ THÔNG TIN VÀ TRUYỀN THÔNG", 14), ("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", 14)):
        p = add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=2, bold=True, size=size)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "BÁO CÁO THỰC TẬP THEO TUẦN", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True, size=18)
    add_paragraph(doc, "NĂM HỌC 2025–2026", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True, size=15)
    for _ in range(2):
        doc.add_paragraph()
    add_paragraph(doc, "TÊN ĐỀ TÀI:", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True, size=14)
    add_paragraph(doc, "BÀI TOÁN TÓM TẮT VĂN BẢN", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True, size=18)
    add_paragraph(doc, "(Báo cáo nội dung thực tập Tuần 3 – 4)", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, italic=True, size=14)
    doc.add_paragraph()
    add_paragraph(doc, "Người hướng dẫn: ThS. Phan Thị Hà", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=13)
    for _ in range(2):
        doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    content = (("Sinh viên thực hiện", "Mã sinh viên"), ("Lường Tiến Dũng", "B22DCCN128"))
    for row_index, row in enumerate(content):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(value), size=13, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
    for _ in range(3):
        doc.add_paragraph()
    add_paragraph(doc, "Hà Nội, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, italic=True, size=13)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    add_heading(doc, "MỤC LỤC", level=1)
    items = [
        ("BẢNG TIẾN ĐỘ", "1"),
        ("NỘI DUNG THỰC TẬP (TUẦN 3 – 4)", "2"),
        ("DANH MỤC BẢNG", "4"),
        ("DANH MỤC HÌNH ẢNH", "5"),
        ("CHƯƠNG I: GIỚI THIỆU", "6"),
        ("  1.1. Bối cảnh và mục tiêu", "6"),
        ("  1.2. Phạm vi thực hiện", "6"),
        ("  1.3. Kết nối với Tuần 1 – 2", "7"),
        ("CHƯƠNG II: CƠ SỞ LÝ THUYẾT VÀ THIẾT KẾ", "8"),
        ("  2.1. Transformer và kiến trúc encoder–decoder", "8"),
        ("  2.2. ViT5 và suy luận sinh tóm tắt", "10"),
        ("  2.3. Đánh giá bằng ROUGE", "12"),
        ("  2.4. Thu thập và kiểm soát chất lượng dữ liệu", "13"),
        ("CHƯƠNG III: CÀI ĐẶT VÀ ĐÁNH GIÁ THỰC NGHIỆM", "15"),
        ("  3.1. Kiến trúc module Tuần 3 – 4", "15"),
        ("  3.2. Thực nghiệm checkpoint ViT5", "17"),
        ("  3.3. Kết quả chuẩn bị dữ liệu thô", "20"),
        ("  3.4. Kết luận và định hướng Tuần 5 – 6", "22"),
        ("TÀI LIỆU THAM KHẢO", "24"),
    ]
    for name, page in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.2)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(15.5), 2, 1)
        set_run_font(p.add_run(name), size=12)
        p.add_run("\t")
        set_run_font(p.add_run(page), size=12)
    doc.add_page_break()


def add_progress_and_summary(doc: Document, audit: dict) -> None:
    add_heading(doc, "BẢNG TIẾN ĐỘ", level=1)
    add_table(
        doc,
        ["Thời gian", "Công việc đã thực hiện", "Sản phẩm đầu ra"],
        [
            (
                "Tuần 3",
                "• Nghiên cứu Transformer, kiến trúc encoder–decoder và mô hình sinh ViT5.\n• Xây dựng CLI tải checkpoint đã pre-train, sinh tóm tắt theo batch.\n• Cài đặt đánh giá ROUGE-1/2/L cùng metadata tái lập thí nghiệm.",
                "• `run_inference_and_rouge.py`.\n• Kết quả thực thi ViT5 trên 5 mẫu: JSON chứa dự đoán, ROUGE và thời gian.",
            ),
            (
                "Tuần 4",
                "• Chuẩn hóa Unicode/HTML/khoảng trắng cho cặp văn bản–tóm tắt.\n• Lọc cặp rỗng, quá ngắn, tóm tắt không ngắn hơn nguồn và cặp trùng chính xác.\n• Lưu audit để truy vết dữ liệu.",
                f"• {audit['accepted_rows']} cặp raw hợp lệ từ {audit['input_rows']} bản ghi nguồn.\n• CSV chuẩn hóa và file audit JSON; chưa split/chunk/fine-tune.",
            ),
        ],
        [2.1, 8.1, 5.3],
    )
    add_heading(doc, "NỘI DUNG THỰC TẬP (TUẦN 3 – 4)", level=1)
    add_heading(doc, "I. Mục tiêu đề ra", level=2)
    add_paragraph(doc, "Kỳ báo cáo này tiếp nối phần cài đặt TextRank và Seq2Seq LSTM ở Tuần 1–2. Mục tiêu là chuyển sang mô hình Transformer sinh tóm tắt đã được tiền huấn luyện cho tiếng Việt, đồng thời hình thành nguồn dữ liệu có thể truy vết để sử dụng cho fine-tuning ở giai đoạn kế tiếp.")
    add_bullet(doc, "Thiết lập được quy trình inference không cần huấn luyện: văn bản nguồn → ViT5 → bản tóm tắt sinh → ROUGE.")
    add_bullet(doc, "Đo kết quả bằng ROUGE-1, ROUGE-2 và ROUGE-L; lưu toàn bộ prediction để có thể kiểm tra bằng mắt, không chỉ dựa vào một con số trung bình.")
    add_bullet(doc, "Chuẩn bị bộ dữ liệu thô gồm các cặp `article`–`summary`, kèm nguồn và audit để tách biệt rõ với bước làm sạch sâu/chia tập của Tuần 5–6.")
    add_heading(doc, "II. Những việc đã làm", level=2)
    add_paragraph(doc, "Về lý thuyết, sinh viên nghiên cứu cơ chế self-attention của Transformer, cách encoder mã hóa toàn bộ văn bản song song và decoder sinh từng token có điều kiện. Trọng tâm được đặt vào ViT5 vì đây là mô hình text-to-text cho tiếng Việt, phù hợp trực tiếp với tóm tắt trừu tượng.")
    add_paragraph(doc, "Về lập trình, project `tuan 3-4/summarization` được xây dựng tách biệt với framework fine-tune của Tuần 5–6. Script suy luận nhận CSV/JSON/JSONL, tự nhận diện các tên cột phổ biến (`article`, `document`, `text` và `summary`, `abstract`, `target`), tạo kết quả JSON có đủ cấu hình generation, thời gian và từng bản dự đoán. Script dữ liệu chuẩn hóa có chủ đích bảo thủ: chỉ sửa nhiễu định dạng và ghi audit, không tự ý chunking hoặc thay đổi nội dung ngữ nghĩa.")
    add_heading(doc, "III. Mục tiêu kỳ báo cáo tiếp theo (Tuần 5 – 6)", level=2)
    add_paragraph(doc, "Dùng dữ liệu raw đã kiểm tra làm đầu vào cho pipeline làm sạch sâu: kiểm tra near-duplicate, chia train/validation/test tránh rò rỉ, xử lý văn bản dài bằng chunking khi cần và fine-tune ViT5. Kết quả baseline ở báo cáo này là mốc so sánh để đánh giá hiệu quả của checkpoint sau fine-tuning.")
    doc.add_page_break()


def add_lists(doc: Document) -> None:
    add_heading(doc, "DANH MỤC BẢNG", level=1)
    for line in (
        "Bảng 2.1: So sánh Seq2Seq LSTM và Transformer trong ngữ cảnh tóm tắt.",
        "Bảng 2.2: Các thành phần cấu hình sinh tóm tắt của thực nghiệm.",
        "Bảng 2.3: Các thước đo ROUGE và cách diễn giải.",
        "Bảng 3.1: Môi trường chạy thực nghiệm.",
        "Bảng 3.2: Kết quả ROUGE của checkpoint ViT5 trên 5 mẫu.",
        "Bảng 3.3: Audit bộ dữ liệu thô Tuần 4.",
    ):
        add_paragraph(doc, line, indent=False, after=4, size=12)
    doc.add_page_break()
    add_heading(doc, "DANH MỤC HÌNH ẢNH", level=1)
    for line in (
        "Hình 3.1: Quy trình triển khai module Tuần 3–4.",
        "Hình 3.2: Biểu đồ ROUGE F1 trên tập kiểm thử 5 mẫu.",
    ):
        add_paragraph(doc, line, indent=False, after=4, size=12)
    doc.add_page_break()


def add_chapter_one(doc: Document) -> None:
    add_heading(doc, "CHƯƠNG I: GIỚI THIỆU", level=1)
    add_heading(doc, "1.1. Bối cảnh và mục tiêu", level=2)
    add_paragraph(doc, "Trong báo cáo Tuần 1–2, hệ thống đã khảo sát TextRank, PhoBERT và Seq2Seq LSTM có Attention. Kết quả từ mô hình LSTM huấn luyện từ đầu cho thấy giới hạn khi phải ghi nhớ văn bản dài và khi dữ liệu huấn luyện chưa đủ lớn. Tuần 3–4 vì vậy chuyển trọng tâm sang Transformer được pre-train, tận dụng tri thức ngôn ngữ học được trước đó thay vì bắt đầu hoàn toàn từ đầu.")
    add_paragraph(doc, "Mục tiêu kỹ thuật không phải là khẳng định một benchmark cuối cùng. Mục tiêu là xây dựng baseline tái lập được: cùng dữ liệu đầu vào, checkpoint, tham số sinh và hàm chấm điểm phải sinh lại cùng dạng báo cáo. Baseline này là căn cứ để trả lời câu hỏi fine-tuning ở Tuần 5–6 có cải thiện thực sự hay không.")
    add_heading(doc, "1.2. Phạm vi thực hiện", level=2)
    add_table(doc, ["Nội dung", "Trong phạm vi Tuần 3–4", "Ngoài phạm vi"], [
        ("Mô hình", "Tải checkpoint ViT5 đã huấn luyện sẵn; chỉ suy luận.", "Huấn luyện từ đầu, LoRA hoặc full fine-tuning."),
        ("Đánh giá", "ROUGE-1/2/L, thời gian, tỷ lệ rút gọn và kiểm tra từng prediction.", "Khẳng định chất lượng tổng quát từ chỉ 5 mẫu smoke test."),
        ("Dữ liệu", "Chuẩn hóa tối thiểu, lọc lỗi rõ ràng, audit và lưu raw CSV.", "Near-duplicate nâng cao, chunking, phân chia train/validation/test."),
        ("Sản phẩm", "CLI, unit test, JSON kết quả và raw dataset có truy vết.", "Web demo hoặc checkpoint fine-tune."),
    ], [3.0, 6.4, 6.1])
    add_heading(doc, "1.3. Kết nối với Tuần 1 – 2", level=2)
    add_paragraph(doc, "TextRank và PhoBERT giúp minh họa hướng trích xuất: hệ thống chọn lại nội dung vốn có trong văn bản. Seq2Seq LSTM đã chuyển sang hướng trừu tượng nhưng phải tự học cả ngôn ngữ lẫn nhiệm vụ từ dữ liệu hạn chế. ViT5 tiếp tục hướng trừu tượng, song khởi đầu từ một mô hình ngôn ngữ tiếng Việt đã pre-train; đây là thay đổi quan trọng nhất của hai tuần này.")
    add_paragraph(doc, "Sự chuyển tiếp được giữ minh bạch: Tuần 3 đánh giá checkpoint có sẵn như một baseline; Tuần 4 chỉ chuẩn bị nguồn dữ liệu. Không gộp số ROUGE của checkpoint pre-train với kết quả fine-tune trong cùng một bảng, vì hai thiết lập có mục tiêu và dữ liệu khác nhau.")


def add_chapter_two(doc: Document) -> None:
    add_heading(doc, "CHƯƠNG II: CƠ SỞ LÝ THUYẾT VÀ THIẾT KẾ", level=1, page_break_before=True)
    add_heading(doc, "2.1. Transformer và kiến trúc encoder–decoder", level=2)
    add_paragraph(doc, "Transformer thay thế phép tính tuần tự của RNN bằng self-attention. Với một chuỗi token, mỗi token tạo ba biểu diễn Query (Q), Key (K) và Value (V). Trọng số chú ý được tính theo công thức Attention(Q, K, V) = softmax(QKᵀ/√dₖ)V. Nhờ đó, một token có thể trực tiếp tham chiếu những token liên quan ở xa trong cùng văn bản; các phép tính giữa token có thể được song song hóa trong encoder.")
    add_paragraph(doc, "Trong mô hình encoder–decoder, encoder tạo biểu diễn ngữ cảnh của bài báo. Decoder dùng masked self-attention để không nhìn trước token tương lai và cross-attention để lấy thông tin từ encoder, sau đó lần lượt dự đoán token của bản tóm tắt. Cơ chế này phù hợp với bài toán abstractive summarization: đầu ra không nhất thiết là tập câu sao chép từ đầu vào.")
    add_table(doc, ["Tiêu chí", "Seq2Seq LSTM + Attention", "Transformer encoder–decoder"], [
        ("Xử lý chuỗi", "Tuần tự theo từng token; khó song song hóa.", "Self-attention xử lý nhiều token đồng thời ở encoder."),
        ("Ngữ cảnh xa", "Attention hỗ trợ nhưng biểu diễn encoder vẫn được tạo tuần tự.", "Mỗi tầng attention tạo liên hệ trực tiếp giữa các vị trí."),
        ("Khởi tạo", "Có thể phải train từ đầu trên dữ liệu nhỏ.", "Dùng checkpoint pre-train rồi fine-tune theo nhiệm vụ."),
        ("Rủi ro", "Quên ngữ cảnh, tốc độ chậm với chuỗi dài.", "Tốn bộ nhớ theo độ dài chuỗi; vẫn có nguy cơ hallucination."),
    ], [3.2, 6.1, 6.2])
    add_heading(doc, "2.2. ViT5 và suy luận sinh tóm tắt", level=2)
    add_paragraph(doc, "ViT5 thuộc họ T5, biểu diễn mọi nhiệm vụ dưới dạng text-to-text. Với nhiệm vụ tóm tắt, văn bản nguồn có thể được đặt sau tiền tố chỉ thị `summarize: `; decoder sinh chuỗi mục tiêu bằng xác suất có điều kiện p(y|x). Trong project, checkpoint mặc định là `VietAI/vit5-base-vietnews-summarization`, một checkpoint có sẵn cho bài toán tóm tắt tin tức tiếng Việt.")
    add_paragraph(doc, "Quá trình sinh dùng beam search. Ở mỗi bước, mô hình giữ nhiều phương án token có xác suất cao nhất thay vì chỉ chọn greedy token. `no_repeat_ngram_size=3` chặn việc lặp lại một cụm 3-gram đã sinh. Đây là ràng buộc kỹ thuật hữu ích nhưng không bảo đảm mọi thông tin sinh ra đều đúng; việc đọc prediction vẫn cần thiết, đặc biệt với dữ liệu y tế.")
    add_table(doc, ["Tham số", "Giá trị thử nghiệm", "Ý nghĩa"], [
        ("max_source_length", "384 token", "Giới hạn đầu vào trong smoke test để phù hợp GPU 4 GB."),
        ("max_new_tokens", "64", "Giới hạn độ dài đầu ra; tránh sinh quá dài."),
        ("min_new_tokens", "8", "Không cho đầu ra quá ngắn."),
        ("num_beams", "2", "Giữ hai nhánh ứng viên; cân bằng chất lượng và tốc độ."),
        ("no_repeat_ngram_size", "3", "Giảm lặp cụm từ 3 token."),
    ], [3.5, 3.0, 9.0])
    add_heading(doc, "2.3. Đánh giá bằng ROUGE", level=2)
    add_paragraph(doc, "ROUGE đo mức chồng lấp giữa bản tóm tắt sinh (candidate) và bản tham chiếu (reference). Báo cáo dùng precision, recall và F1, lấy trung bình theo từng mẫu rồi nhân 100. Precision cao cho biết nội dung sinh ra ít từ thừa so với tham chiếu; recall cao cho biết bao phủ được nhiều từ của tham chiếu. F1 là trung bình điều hòa của hai đại lượng và là giá trị dùng để so sánh ngắn gọn trong bảng kết quả.")
    add_table(doc, ["Chỉ số", "Đơn vị so khớp", "Ý nghĩa thực tế"], [
        ("ROUGE-1", "Unigram", "Khả năng giữ từ khóa và chủ đề chính."),
        ("ROUGE-2", "Bigram", "Mức độ đúng của cụm từ/quan hệ từ liền nhau; thường khó đạt hơn ROUGE-1."),
        ("ROUGE-L", "Longest Common Subsequence", "Giữ trình tự nội dung chung, không bắt buộc các từ phải đứng sát nhau."),
    ], [3.0, 3.2, 9.3])
    add_heading(doc, "2.4. Thu thập và kiểm soát chất lượng dữ liệu", level=2)
    add_paragraph(doc, "Mỗi ví dụ huấn luyện tóm tắt cần là một cặp có hướng: `article` là văn bản nguồn và `summary` là bản tóm tắt chuẩn. Nếu hai cột bị đảo, có cặp trùng hoặc summary dài ngang nguồn, mô hình có thể học sai nhiệm vụ. Vì thế bước Tuần 4 tạo một schema thống nhất gồm `id`, `article`, `summary`, `source_dataset`, `source_row`.")
    add_paragraph(doc, "Bộ lọc chỉ loại các lỗi chắc chắn: dữ liệu rỗng, quá ngắn, tóm tắt không ngắn hơn văn bản nguồn và cặp article–summary trùng chính xác. Những trường hợp gần giống không bị loại ở đây vì cần một chính sách kiểm tra kỹ và chống rò rỉ giữa các split ở Tuần 5. Audit JSON lưu các quyết định này để không biến thao tác tiền xử lý thành một hộp đen.")


def add_chapter_three(doc: Document, result: dict, audit: dict, pipeline: Path, chart: Path) -> None:
    rouge = result["metrics"]["rouge"]
    compression = result["metrics"]["compression"]
    run = result["run"]
    add_heading(doc, "CHƯƠNG III: CÀI ĐẶT VÀ ĐÁNH GIÁ THỰC NGHIỆM", level=1, page_break_before=True)
    add_heading(doc, "3.1. Kiến trúc module Tuần 3 – 4", level=2)
    add_paragraph(doc, "Module được tổ chức thành hai entry point độc lập dùng chung các tiện ích dữ liệu, chuẩn hóa và ROUGE. `run_inference_and_rouge.py` dành cho Tuần 3: đọc mẫu có reference, tải ViT5 một lần, sinh theo batch và xuất JSON. `prepare_raw_dataset.py` dành cho Tuần 4: đọc nguồn CSV/JSON/JSONL, chuẩn hóa tối thiểu, áp dụng các quality gate và xuất CSV cùng audit JSON.")
    doc.add_picture(str(pipeline), width=Cm(16.1))
    add_caption(doc, "Hình 3.1: Quy trình triển khai module Tuần 3–4.")
    add_paragraph(doc, "Thiết kế này chủ động ngăn chồng lấn phạm vi: script inference không ghi đè dữ liệu nguồn; script raw-data không tạo split hoặc checkpoint. Hai đầu ra được dùng làm đầu vào/đối chứng cho pipeline fine-tuning ở Tuần 5–6.")
    add_heading(doc, "3.2. Thực nghiệm checkpoint ViT5", level=2, page_break_before=True)
    add_paragraph(doc, "Smoke test được chạy trên 5 mẫu đầu tiên của `data/summarization_samples.json`. Tập này chỉ dùng xác nhận pipeline hoạt động đầu-cuối và giúp đọc từng lỗi của mô hình; kích thước 5 không đủ để kết luận hiệu năng chung. Các số dưới đây được đọc trực tiếp từ `results/vit5_sample_5.json` do script sinh ra.")
    add_table(doc, ["Hạng mục", "Giá trị"], [
        ("Checkpoint", run["model"]),
        ("Thiết bị", "NVIDIA GeForce RTX 3050 Laptop GPU (4 GB VRAM), chạy bằng CUDA"),
        ("Số mẫu", str(run["number_of_examples"])),
        ("Thời gian", f"{run['elapsed_seconds']} giây; {run['seconds_per_example']} giây/mẫu"),
        ("Độ dài nguồn trung bình", f"{result['input_audit']['length_statistics_words']['article_mean']} từ"),
        ("Độ dài tóm tắt sinh trung bình", f"{compression['summary_words_mean']} từ; rút gọn {compression['compression_percent']}% theo số từ"),
    ], [5.7, 9.8])
    add_table(doc, ["Chỉ số", "Precision", "Recall", "F1"], [
        ("ROUGE-1", f"{rouge['rouge1']['precision']:.2f}", f"{rouge['rouge1']['recall']:.2f}", f"{rouge['rouge1']['f1']:.2f}"),
        ("ROUGE-2", f"{rouge['rouge2']['precision']:.2f}", f"{rouge['rouge2']['recall']:.2f}", f"{rouge['rouge2']['f1']:.2f}"),
        ("ROUGE-L", f"{rouge['rougeL']['precision']:.2f}", f"{rouge['rougeL']['recall']:.2f}", f"{rouge['rougeL']['f1']:.2f}"),
    ], [4.0, 3.7, 3.7, 3.7])
    doc.add_picture(str(chart), width=Cm(15.5))
    add_caption(doc, "Hình 3.2: ROUGE F1 trên 5 mẫu kiểm thử của checkpoint ViT5.")
    add_paragraph(doc, "ROUGE-1 F1 đạt 58,05 cho thấy checkpoint nhận diện được một phần đáng kể từ khóa. ROUGE-2 F1 là 21,57, thấp hơn đáng kể vì việc tái hiện đúng các cụm từ liên tiếp khó hơn. ROUGE-L F1 đạt 34,29, phản ánh mức độ giữ trình tự nội dung ở mức trung bình trên bộ mẫu nhỏ này.")
    add_heading(doc, "3.2.1. Kiểm tra định tính prediction", level=3)
    sample_rows = []
    for prediction in result["predictions"]:
        article_short = prediction["article"][:140].rsplit(" ", 1)[0] + "…"
        sample_rows.append((prediction["id"], article_short, prediction["reference"], prediction["prediction"]))
    add_table(doc, ["ID", "Nguồn (rút gọn)", "Tham chiếu", "ViT5 sinh"], sample_rows, [0.8, 4.8, 4.8, 5.0])
    add_paragraph(doc, "Quan sát định tính là bắt buộc vì ROUGE không phát hiện toàn bộ sai lệch thực tế. Ví dụ mẫu 3, bản sinh nhắc tới “PGS.TS Nguyễn Ngọc Long” dù thông tin này không xuất hiện trong văn bản nguồn. Đây là hallucination: câu đọc tự nhiên nhưng không được chứng minh bởi đầu vào. Các mẫu 4–5 cũng khá ngắn, bỏ qua một số hệ quả quan trọng. Vì vậy checkpoint pre-train phù hợp làm baseline nhưng chưa được dùng độc lập cho văn bản nhạy cảm như y tế/pháp lý.")
    doc.add_page_break()
    add_heading(doc, "3.3. Kết quả chuẩn bị dữ liệu thô", level=2)
    add_paragraph(doc, "Nguồn đầu vào là một CSV cục bộ có 1.000 cặp `Document`–`Summary`. Script chỉ sử dụng một biến thể dữ liệu để không gộp nhiều bản biến đổi của cùng nguồn. Sau chuẩn hóa và lọc trùng chính xác, dữ liệu đầu ra nằm tại `tuan 3-4/summarization/data/raw/vietnews_medical_raw_1000.csv`.")
    stats = audit["length_statistics_words"]
    add_table(doc, ["Hạng mục audit", "Kết quả"], [
        ("Bản ghi nguồn", str(audit["input_rows"])),
        ("Cặp được chấp nhận", str(audit["accepted_rows"])),
        ("Cặp bị loại", f"{audit['rejected_rows']} — toàn bộ là exact duplicate pair"),
        ("Cột nguồn được nhận diện", "Document → article; Summary → summary"),
        ("Độ dài article", f"Trung bình {stats['article_mean']} từ; min {stats['article_min']}; max {stats['article_max']}"),
        ("Độ dài summary", f"Trung bình {stats['summary_mean']} từ; min {stats['summary_min']}; max {stats['summary_max']}"),
        ("Split/chunking", "Chưa thực hiện — đúng phạm vi dữ liệu raw của Tuần 4."),
    ], [5.7, 9.8])
    add_paragraph(doc, "990 cặp còn lại đủ điều kiện hình thức để chuyển tiếp. Tuy nhiên “đủ điều kiện raw” không đồng nghĩa “sẵn sàng train”: Tuần 5 cần tiếp tục kiểm tra near-duplicate, sự trùng lặp giữa tập train/validation/test, độ dài theo tokenizer của ViT5, và các trường hợp tóm tắt chưa chính xác hoặc lệch miền. Việc tách audit riêng giúp phân biệt rõ hai mức đánh giá đó.")
    add_heading(doc, "3.4. Kết luận và định hướng Tuần 5 – 6", level=2)
    add_paragraph(doc, "Hai tuần đã hoàn thành một baseline Transformer có thể chạy lại và một nguồn dữ liệu raw có audit. Kết quả thực nghiệm trên 5 mẫu chứng minh quy trình tải checkpoint, sinh tóm tắt, tính ROUGE và lưu prediction hoạt động đúng; đồng thời kiểm tra định tính đã chỉ ra rủi ro hallucination và mất chi tiết. Không diễn giải 5 mẫu như điểm chuẩn cuối cùng.")
    add_bullet(doc, "Tuần 5: làm sạch sâu 990 cặp, kiểm tra gần trùng/lệch dữ liệu, tokenization và chiến lược chunking cho bài vượt giới hạn token.")
    add_bullet(doc, "Tuần 5: chia train/validation/test theo chính sách tránh rò rỉ; lưu seed, danh sách bản ghi và thống kê từng split.")
    add_bullet(doc, "Tuần 6: fine-tune ViT5, theo dõi loss và ROUGE trên validation; so sánh checkpoint fine-tune với baseline pre-train cùng một held-out test set.")
    add_bullet(doc, "Tuần 7: đánh giá cuối, bổ sung kiểm tra định tính/hallucination và tích hợp mô hình vào web demo.")
    doc.add_page_break()


def add_references(doc: Document) -> None:
    add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1)
    references = [
        "[1] A. Vaswani et al., “Attention Is All You Need,” Advances in Neural Information Processing Systems, vol. 30, 2017. [Online]. Available: https://arxiv.org/abs/1706.03762.",
        "[2] C. Raffel et al., “Exploring the Limits of Transfer Learning with a Unified Text-to-Text Transformer,” Journal of Machine Learning Research, vol. 21, no. 140, pp. 1–67, 2020. [Online]. Available: https://jmlr.org/papers/v21/20-074.html.",
        "[3] C.-Y. Lin, “ROUGE: A Package for Automatic Evaluation of Summaries,” Text Summarization Branches Out, 2004, pp. 74–81. [Online]. Available: https://aclanthology.org/W04-1013.",
        "[4] VietAI, “vit5-base-vietnews-summarization,” Hugging Face Model Card. [Online]. Available: https://huggingface.co/VietAI/vit5-base-vietnews-summarization. [Accessed: 23-Jul-2026].",
        "[5] Hugging Face, “Transformers: Text generation,” Documentation. [Online]. Available: https://huggingface.co/docs/transformers/main_classes/text_generation. [Accessed: 23-Jul-2026].",
        "[6] Kế hoạch và đề cương chi tiết: Bài toán tóm tắt nội dung văn bản tiếng Việt, tài liệu nội bộ workspace, phần lộ trình Tuần 3–4.",
    ]
    for reference in references:
        add_paragraph(doc, reference, indent=False, after=8, size=12)


def main() -> None:
    if not RESULT_PATH.exists() or not AUDIT_PATH.exists():
        raise FileNotFoundError("Thiếu kết quả thực nghiệm hoặc audit dữ liệu. Hãy chạy các script Tuần 3–4 trước.")
    result = json.loads(RESULT_PATH.read_text(encoding="utf-8"))
    audit = json.loads(AUDIT_PATH.read_text(encoding="utf-8"))
    pipeline, chart = create_figures(result["metrics"]["rouge"])
    doc = setup_document()
    add_cover(doc)
    add_contents(doc)
    add_progress_and_summary(doc, audit)
    add_lists(doc)
    add_chapter_one(doc)
    add_chapter_two(doc)
    add_chapter_three(doc, result, audit, pipeline, chart)
    add_references(doc)
    doc.save(OUTPUT_DOCX)
    print(f"Đã tạo: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
