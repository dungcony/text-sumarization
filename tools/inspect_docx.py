#!/usr/bin/env python3
"""Print paragraph, table, section, and image metadata from a DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--start", type=int, default=0)
    parser.add_argument("--end", type=int)
    args = parser.parse_args()

    document = Document(args.docx)
    paragraphs = document.paragraphs
    end = len(paragraphs) if args.end is None else min(args.end, len(paragraphs))

    print(f"paragraphs={len(paragraphs)} tables={len(document.tables)} sections={len(document.sections)}")
    for index, paragraph in enumerate(paragraphs[args.start:end], start=args.start):
        text = paragraph.text.replace("\t", "\\t").replace("\n", "\\n")
        drawing_count = len(paragraph._p.xpath(".//w:drawing"))
        page_break = bool(paragraph._p.xpath(".//w:br[@w:type='page']"))
        section_break = paragraph._p.pPr is not None and paragraph._p.pPr.sectPr is not None
        num_id = ""
        ilvl = ""
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            num_pr = paragraph._p.pPr.numPr
            if num_pr.numId is not None:
                num_id = num_pr.numId.val
            if num_pr.ilvl is not None:
                ilvl = num_pr.ilvl.val
        print(
            f"P{index:03d} style={paragraph.style.name!r} num={num_id}/{ilvl} "
            f"draw={drawing_count} pagebr={page_break} sectbr={section_break} text={text!r}"
        )
        for drawing in paragraph._p.xpath(".//w:drawing"):
            blips = drawing.xpath(".//a:blip")
            extents = drawing.xpath(".//wp:extent")
            rel_id = blips[0].get(qn("r:embed")) if blips else ""
            extent = ""
            if extents:
                extent = f"{extents[0].get('cx')}x{extents[0].get('cy')}"
            print(f"    drawing rel={rel_id} extent={extent}")

    for index, table in enumerate(document.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        preview = " | ".join(cell.text.replace("\n", " / ") for cell in table.rows[0].cells)
        print(f"T{index:02d} rows={rows} cols={cols} first_row={preview!r}")

    for index, section in enumerate(document.sections):
        print(
            f"S{index:02d} size={section.page_width}x{section.page_height} "
            f"margins={section.left_margin},{section.right_margin},"
            f"{section.top_margin},{section.bottom_margin}"
        )

    rel_images = [
        rel for rel in document.part.rels.values()
        if rel.reltype.endswith("/image")
    ]
    print(f"image_relationships={len(rel_images)}")
    for rel in rel_images:
        target = getattr(rel.target_part, "partname", "")
        print(f"  {rel.rId}: {target}")

    body = document._element.body
    child_names = [child.tag.rsplit("}", 1)[-1] for child in body]
    print(f"body_children={len(child_names)} p={child_names.count('p')} tbl={child_names.count('tbl')}")


if __name__ == "__main__":
    main()
