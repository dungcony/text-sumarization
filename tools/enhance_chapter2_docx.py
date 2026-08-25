#!/usr/bin/env python3
"""Rewrite Chapter 1 and expand Chapter 2 with accessible technical figures."""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor


FIGURE_DIR = Path("bao-cao/figures_chapter2")
BODY_FONT = "Times New Roman"


def first_text_run(paragraph: Paragraph):
    return next((run for run in paragraph.runs if run.text), paragraph.runs[0] if paragraph.runs else None)


def copy_paragraph_properties(source: Paragraph, target: Paragraph) -> None:
    for ppr in target._p.xpath("./w:pPr"):
        target._p.remove(ppr)
    if source._p.pPr is not None:
        target._p.insert(0, deepcopy(source._p.pPr))


def copy_run_properties(source_run, target_run) -> None:
    for rpr in target_run._r.xpath("./w:rPr"):
        target_run._r.remove(rpr)
    if source_run is not None and source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def set_font(
    run,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
    font_name: str = BODY_FONT,
) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def insert_paragraph_after(anchor: Paragraph, template: Paragraph) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    copy_paragraph_properties(template, paragraph)
    return paragraph


def add_text_after(
    anchor: Paragraph,
    template: Paragraph,
    text: str,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, template)
    run = paragraph.add_run(text)
    copy_run_properties(first_text_run(template), run)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return paragraph


def replace_text(paragraph: Paragraph, text: str) -> None:
    template_run = first_text_run(paragraph)
    run_properties = deepcopy(template_run._r.rPr) if template_run is not None and template_run._r.rPr is not None else None
    paragraph.clear()
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def add_callout_after(
    anchor: Paragraph,
    template: Paragraph,
    label: str,
    body: str,
    *,
    fill: str = "EAF3F8",
    accent: str = "1F8A70",
) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, template)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.keep_together = True

    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")

    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    pbdr.append(left)

    label_run = paragraph.add_run(label)
    copy_run_properties(first_text_run(template), label_run)
    set_font(label_run, 12, bold=True, color=accent)
    body_run = paragraph.add_run(body)
    copy_run_properties(first_text_run(template), body_run)
    set_font(body_run, 12)
    return paragraph


def add_code_after(
    anchor: Paragraph,
    template: Paragraph,
    title: str,
    code: str,
    source_number: int,
) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, template)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True

    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F8")
    shd.set(qn("w:val"), "clear")
    ppr.append(shd)

    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "4C78A8")
    pbdr.append(left)
    ppr.append(pbdr)

    title_run = paragraph.add_run(f"{title} [{source_number}]\n")
    set_font(title_run, 10.5, bold=True, color="234E70")
    code_run = paragraph.add_run(code)
    set_font(code_run, 9.5, font_name="Consolas")
    note_run = paragraph.add_run(
        "\nMã chỉ minh họa cách gọi API công khai; dữ liệu và tham số không phải cấu hình của đề tài."
    )
    set_font(note_run, 9, italic=True, color="5B6573")
    return paragraph


def set_image_alt(paragraph: Paragraph, title: str, description: str) -> None:
    doc_props = paragraph._p.xpath(".//wp:docPr")
    if not doc_props:
        return
    doc_prop = doc_props[-1]
    doc_prop.set("title", title)
    doc_prop.set("descr", description)


def put_picture(paragraph: Paragraph, path: Path, width: float, title: str, description: str) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    set_image_alt(paragraph, title, description)


def add_picture_after(
    anchor: Paragraph,
    image_template: Paragraph,
    path: Path,
    width: float,
    title: str,
    description: str,
) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, image_template)
    put_picture(paragraph, path, width, title, description)
    return paragraph


def add_caption_after(anchor: Paragraph, caption_template: Paragraph, text: str) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, caption_template)
    run = paragraph.add_run(text)
    copy_run_properties(first_text_run(caption_template), run)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = False
    return paragraph


def add_bullet_after(anchor: Paragraph, bullet_template: Paragraph, text: str) -> Paragraph:
    return add_text_after(anchor, bullet_template, text)


def clear_separator(paragraph: Paragraph) -> None:
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def format_equation(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    for run in paragraph.runs:
        set_font(run, 13)


def format_reference_entry(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.first_line_indent = Inches(-0.28)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True
    for run in paragraph.runs:
        set_font(run, 11, bold=False, italic=False)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table_width = sum(widths)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for grid_col in list(grid):
        grid.remove(grid_col)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")

            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in (("top", 90), ("left", 110), ("bottom", 90), ("right", 110)):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def replace_literature_table(document: Document, old_table, caption: Paragraph) -> None:
    headers = ["STT", "Công trình", "Năm", "Hướng tiếp cận", "Đóng góp và giới hạn"]
    rows = [
        ("1", "Luhn [1]", "1958", "Trích xuất thống kê", "Dùng tần suất từ để nhận diện nội dung nổi bật; chưa mô hình hóa ngữ cảnh."),
        ("2", "Edmundson [2]", "1969", "Đặc trưng thủ công", "Bổ sung từ chỉ dẫn, tiêu đề và vị trí câu; phụ thuộc vào luật thiết kế sẵn."),
        ("3", "Mihalcea & Tarau [3]", "2004", "Xếp hạng đồ thị", "TextRank không cần dữ liệu gán nhãn; đầu ra vẫn là các đơn vị trích xuất."),
        ("4", "Rush và cộng sự [4]", "2015", "Sinh tóm tắt với Attention", "Đưa học sâu sinh tạo vào tóm tắt; nghiên cứu ban đầu tập trung ở mức câu."),
        ("5", "See và cộng sự [5]", "2017", "Pointer-Generator", "Kết hợp sao chép và sinh từ, dùng coverage để giảm lặp; không loại bỏ hoàn toàn sai lệch."),
        ("6", "Lewis và cộng sự [7]", "2020", "BART tiền huấn luyện", "Mạnh cho các tác vụ sinh văn bản; cần dữ liệu và tài nguyên tính toán đáng kể."),
        ("7", "Nguyễn & Nguyễn [11]", "2020", "PhoBERT", "Cung cấp biểu diễn ngữ cảnh cho tiếng Việt; bản thân PhoBERT không phải mô hình sinh tóm tắt."),
        ("8", "Phan và cộng sự [12]", "2022", "ViT5 Encoder-Decoder", "Hỗ trợ sinh văn bản và tóm tắt tiếng Việt; chất lượng phụ thuộc dữ liệu, độ dài ngữ cảnh và giải mã."),
    ]

    table = document.add_table(rows=1, cols=len(headers))
    table.style = old_table.style
    table._tbl.getparent().remove(table._tbl)
    caption._p.addnext(table._tbl)

    header_row = table.rows[0]
    tr_pr = header_row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)

    for index, text in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "234E70")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_font(run, 9, bold=True, color="FFFFFF")

    for row_index, row_data in enumerate(rows, start=1):
        row = table.add_row()
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)
        for column_index, text in enumerate(row_data):
            cell = row.cells[column_index]
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 0:
                shade_cell(cell, "F4F6F8")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if column_index in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_font(run, 8.5)

    set_table_geometry(table, [600, 1900, 700, 2050, 4110])
    old_table._tbl.getparent().remove(old_table._tbl)


def enhance(input_path: Path, output_path: Path) -> None:
    document = Document(input_path)
    paragraphs = document.paragraphs
    if len(paragraphs) < 240 or not paragraphs[115].text.startswith("CHƯƠNG 2"):
        raise RuntimeError("Không nhận diện được cấu trúc Chương 2 của tài liệu đầu vào.")

    normal = paragraphs[118]
    heading3 = paragraphs[116]
    heading4 = paragraphs[117]
    bullet = paragraphs[121]
    image_template = paragraphs[135]
    caption_template = paragraphs[136]

    chapter1_heading = paragraphs[47]
    p48, p49, p50, p51 = paragraphs[48:52]
    p52, p53, p54, p55, p56, p57, p58, p59, p60, p61, p62, p63, p64, p65, p66, p67, p68, p69, p70, p71 = paragraphs[52:72]
    p72, p73, p74, p75, p76, p77, p78, p79, p80, p81, p82, p83, p84 = paragraphs[72:85]
    p85, p86, p87, p88, p89, p90, p91 = paragraphs[85:92]
    p92, p93, p94, p95, p96, p97 = paragraphs[92:98]
    p98, p99, p100, p101, p102, p103, p104, p105, p106, p107, p108, p109 = paragraphs[98:110]
    p110, p111, p112, p113, p114 = paragraphs[110:115]
    chapter2_heading = paragraphs[115]
    literature_table = document.tables[0]

    p118 = paragraphs[118]
    p123, p124, p125, p126 = paragraphs[123:127]
    p127, p128, p129, p134 = paragraphs[127], paragraphs[128], paragraphs[129], paragraphs[134]
    p135, p136 = paragraphs[135], paragraphs[136]
    p141, p143, p144 = paragraphs[141], paragraphs[143], paragraphs[144]
    p147, p158, p160, p165 = paragraphs[147], paragraphs[158], paragraphs[160], paragraphs[165]
    p167, p168, p170, p171, p172, p173 = paragraphs[167], paragraphs[168], paragraphs[170], paragraphs[171], paragraphs[172], paragraphs[173]
    p176, p182, p183, p184, p189, p191, p192, p193 = paragraphs[176], paragraphs[182], paragraphs[183], paragraphs[184], paragraphs[189], paragraphs[191], paragraphs[192], paragraphs[193]
    p196, p198, p206, p208, p212 = paragraphs[196], paragraphs[198], paragraphs[206], paragraphs[208], paragraphs[212]
    p216, p217, p218, p219, p221, p222, p226 = paragraphs[216], paragraphs[217], paragraphs[218], paragraphs[219], paragraphs[221], paragraphs[222], paragraphs[226]
    p229, p236 = paragraphs[229], paragraphs[236]
    page_break_before_chapter3, chapter3_heading = paragraphs[238], paragraphs[239]
    post_results_blank, page_break_before_conclusion, conclusion_heading = paragraphs[277], paragraphs[278], paragraphs[279]
    reference_paragraphs = paragraphs[293:301]
    reference_anchor = reference_paragraphs[-1]
    separators = [paragraphs[index] for index in (145, 174, 194, 213, 227, 237)]

    replace_text(chapter1_heading, "CHƯƠNG 1. TỔNG QUAN VỀ BÀI TOÁN TÓM TẮT VĂN BẢN")
    replace_text(p48, "1.1. Bối cảnh và phát biểu bài toán")
    replace_text(
        p49,
        "Các nghiên cứu về tóm tắt tự động xuất hiện từ giai đoạn đầu của xử lý thông tin bằng máy tính. Luhn (1958) đề xuất nhận diện câu quan trọng dựa trên tần suất từ [1]; Edmundson (1969) mở rộng hướng này bằng các tín hiệu như từ chỉ dẫn, tiêu đề và vị trí câu [2]. Những công trình nền tảng đó hình thành hướng tóm tắt trích xuất, trong đó hệ thống lựa chọn các phần đã có trong văn bản nguồn.",
    )
    replace_text(
        p50,
        "Trong môi trường số, một người có thể phải đọc đồng thời tin tức, báo cáo, tài liệu chuyên môn và nội dung trên mạng xã hội. Tóm tắt văn bản tự động nhằm rút gọn lượng thông tin cần đọc bằng cách tạo ra một văn bản ngắn hơn nhưng vẫn giữ các ý quan trọng, mối quan hệ chính và những dữ kiện cần thiết của nguồn.",
    )
    replace_text(
        p51,
        "Bài toán của đề tài được xác định như sau: từ một văn bản tiếng Việt đầu vào, hệ thống tạo ra bản tóm tắt có độ dài phù hợp, bao quát nội dung cốt lõi, hạn chế lặp ý, dễ đọc và không đưa thêm thông tin trái với văn bản nguồn. Đây là bài toán nhiều mục tiêu; một bản tóm tắt ngắn hơn chưa chắc tốt hơn nếu làm mất dữ kiện quan trọng hoặc thay đổi ý nghĩa ban đầu.",
    )
    add_callout_after(
        p51,
        normal,
        "Ví dụ ngắn. ",
        "Văn bản nguồn: “Bệnh viện A khai trương khu cấp cứu 30 giường. Khu này hoạt động liên tục và được bổ sung 12 bác sĩ.” Bản tóm tắt phù hợp có thể là: “Bệnh viện A đưa vào hoạt động khu cấp cứu 30 giường, có thêm 12 bác sĩ.” Câu tóm tắt ngắn hơn nhưng vẫn giữ chủ thể, sự kiện và hai dữ kiện chính.",
        fill="E9F6F2",
    )

    replace_text(p52, "1.2. Hình thức hóa và các yêu cầu của bản tóm tắt")
    replace_text(
        p53,
        "Giả sử văn bản đầu vào gồm n câu theo đúng thứ tự xuất hiện. Cần giữ thông tin về thứ tự này vì một bản tóm tắt trích xuất thường chọn câu theo điểm quan trọng rồi sắp xếp lại theo trật tự của văn bản nguồn để bảo đảm mạch trình bày.",
    )
    replace_text(p54, "Văn bản đầu vào D được biểu diễn bằng một dãy có thứ tự:")
    replace_text(p55, "D = (s₁, s₂, ..., sₙ)")
    format_equation(p55)
    replace_text(
        p56,
        "Mục tiêu của hệ thống là xây dựng một hàm f tạo bản tóm tắt S gồm m đơn vị đầu ra, với m nhỏ hơn n trong trường hợp tóm tắt ở mức câu:",
    )
    replace_text(p57, "S = f(D; θ, B) = (y₁, y₂, ..., yₘ)")
    format_equation(p57)
    replace_text(
        p58,
        "Trong đó, θ là tập tham số hoặc quy tắc của phương pháp; B là ngân sách độ dài. Với tóm tắt trích xuất, mỗi yᵢ là một câu hoặc đoạn được chọn từ D. Với tóm tắt trừu tượng, yᵢ có thể là từ hoặc câu mới do mô hình sinh ra dựa trên nội dung của D.",
    )
    replace_text(
        p59,
        "Không tồn tại một bản tóm tắt tối ưu duy nhất cho mọi người đọc và mọi mục đích. Cùng một văn bản, người đọc tin nhanh có thể cần vài câu kết luận, trong khi người đọc chuyên môn cần giữ phương pháp, số liệu và điều kiện áp dụng. Vì vậy, chất lượng phải được xem xét theo nhiều tiêu chí thay vì chỉ dựa vào tỷ lệ rút gọn.",
    )
    replace_text(p60, "Bốn nhóm yêu cầu chính của bài toán gồm:")
    replace_text(p61, "1. Độ phủ thông tin (Coverage)")
    replace_text(p62, "max Coverage(S, D)")
    format_equation(p62)
    replace_text(
        p63,
        "Độ phủ phản ánh mức độ bản tóm tắt giữ được các ý và dữ kiện quan trọng của nguồn. Một bản tóm tắt có độ phủ thấp có thể đúng ngữ pháp nhưng bỏ mất nguyên nhân, kết quả hoặc con số cần thiết.",
    )
    replace_text(p64, "2. Hạn chế trùng lặp (Redundancy)")
    replace_text(p65, "min Redundancy(S)")
    format_equation(p65)
    replace_text(
        p66,
        "Các câu gần nghĩa không nên cùng chiếm nhiều chỗ trong bản tóm tắt. Hạn chế trùng lặp giúp ngân sách độ dài được dành cho các nhóm ý khác nhau, nhưng việc loại lặp không được làm mất các chi tiết bổ sung thực sự cần thiết.",
    )
    replace_text(p67, "3. Tuân thủ ngân sách độ dài (Length Budget)")
    replace_text(p68, "Length(S) ≤ B")
    format_equation(p68)
    replace_text(
        p69,
        "B có thể được quy định bằng số câu, số từ hoặc số token. Giá trị phù hợp phụ thuộc loại tài liệu và mục đích sử dụng; không có một tỷ lệ cố định áp dụng tốt cho mọi văn bản.",
    )
    replace_text(p70, "4. Mạch lạc và nhất quán với nguồn")
    replace_text(
        p71,
        "Bản tóm tắt phải dễ đọc, giữ được quan hệ giữa chủ thể, hành động, thời gian và số liệu. Phương pháp trích xuất giảm nguy cơ tự tạo dữ kiện nhưng vẫn có thể gây hiểu sai khi lấy câu ra khỏi ngữ cảnh. Phương pháp sinh tạo diễn đạt linh hoạt hơn nhưng có thể sinh thông tin không được nguồn hỗ trợ, do đó cần đánh giá cả tính đúng sự thật.",
    )
    for requirement_heading in (p61, p64, p67, p70):
        if requirement_heading.runs:
            requirement_heading.runs[0].bold = True

    replace_text(p72, "1.3. Các hướng phân loại bài toán tóm tắt")
    replace_text(p73, "1.3.1. Phân loại theo cách tạo ra bản tóm tắt")
    replace_text(p74, "a) Tóm tắt trích xuất (Extractive Summarization)")
    replace_text(
        p75,
        "Phương pháp trích xuất chấm điểm hoặc xếp hạng các câu, sau đó chọn một số câu nguyên bản để tạo bản tóm tắt. Những câu được chọn thường được đưa về thứ tự xuất hiện ban đầu để giữ dòng lập luận của văn bản.",
    )
    replace_text(
        p76,
        "Quy trình phổ biến gồm bốn bước: chia văn bản thành câu; biểu diễn mỗi câu bằng đặc trưng thống kê hoặc vector ngữ nghĩa; ước lượng mức quan trọng và mức trùng lặp; chọn các câu phù hợp với ngân sách độ dài. TextRank là một ví dụ tiêu biểu của hướng này [3].",
    )
    replace_text(
        p77,
        "Ưu điểm của tóm tắt trích xuất là có thể truy vết từng câu về văn bản nguồn, thường không cần mô hình sinh lớn và ít có khả năng tự bịa ra từ hoặc con số mới. Vì thế, phương pháp phù hợp làm đường cơ sở và phù hợp khi người dùng cần kiểm tra nguồn nhanh.",
    )
    replace_text(
        p78,
        "Hạn chế là các câu được ghép từ nhiều vị trí có thể thiếu liên kết, chứa đại từ không rõ đối tượng hoặc lặp lại cấu trúc. Phương pháp cũng khó kết hợp thông tin nằm rải rác thành một câu cô đọng mới.",
    )
    replace_text(p79, "b) Tóm tắt trừu tượng (Abstractive Summarization)")
    replace_text(
        p80,
        "Phương pháp trừu tượng học cách mã hóa nội dung nguồn và sinh ra chuỗi từ mới. Bản tóm tắt có thể diễn đạt lại, gộp nhiều câu và lược bỏ các chi tiết phụ thay vì chỉ sao chép câu nguyên bản.",
    )
    replace_text(
        p81,
        "Ưu điểm là đầu ra có khả năng ngắn gọn và liền mạch hơn, đặc biệt khi nội dung chính nằm ở nhiều câu khác nhau. Các mô hình Attention, Pointer-Generator và Transformer là những cột mốc quan trọng của hướng tiếp cận này [4]-[6].",
    )
    replace_text(
        p82,
        "Hạn chế chính là nhu cầu dữ liệu và tài nguyên tính toán cao hơn, kết quả khó giải thích hơn và có nguy cơ sinh chi tiết không có trong nguồn. Mức độ rủi ro phụ thuộc mô hình, dữ liệu huấn luyện, miền văn bản và cách giải mã.",
    )
    replace_text(p83, "c) Tóm tắt lai (Hybrid Summarization)")
    replace_text(
        p84,
        "Hướng lai kết hợp nhiều cơ chế, chẳng hạn chọn trước các câu hoặc đoạn liên quan rồi dùng mô hình sinh để diễn đạt lại. Cách kết hợp phải được thiết kế theo dữ liệu và mục tiêu cụ thể; không có một tỷ lệ trích xuất/sinh tạo cố định bảo đảm tốt cho mọi trường hợp.",
    )
    for method_heading in (p74, p79, p83):
        if method_heading.runs:
            method_heading.runs[0].bold = True

    replace_text(p85, "1.3.2. Phân loại theo số lượng tài liệu đầu vào")
    replace_text(
        p86,
        "Tóm tắt đơn văn bản (Single-document Summarization) xử lý một tài liệu tại mỗi lần. Thách thức chính là xác định cấu trúc và nội dung quan trọng bên trong tài liệu đó.",
    )
    replace_text(
        p87,
        "Tóm tắt đa văn bản (Multi-document Summarization) tổng hợp thông tin từ nhiều tài liệu cùng chủ đề. Ngoài việc chọn ý quan trọng, hệ thống còn phải xử lý nội dung trùng lặp, khác biệt góc nhìn, mâu thuẫn dữ kiện và thứ tự thời gian giữa các nguồn.",
    )
    replace_text(p88, "1.3.3. Phân loại theo mục tiêu sử dụng")
    replace_text(
        p89,
        "Tóm tắt tổng quát cố gắng phản ánh các ý chính của toàn văn bản mà không ưu tiên một câu hỏi cụ thể. Đây là dạng thường gặp trong tóm tắt tin tức hoặc báo cáo.",
    )
    replace_text(
        p90,
        "Tóm tắt hướng truy vấn chỉ ưu tiên phần nội dung liên quan tới nhu cầu của người dùng, chẳng hạn “kết quả tài chính” hoặc “tác dụng phụ”. Vì vậy, cùng một văn bản có thể tạo ra các bản tóm tắt khác nhau theo truy vấn.",
    )
    replace_text(
        p91,
        "Tóm tắt theo lĩnh vực được điều chỉnh cho đặc trưng của văn bản y tế, pháp luật, khoa học hoặc tài chính. Các lĩnh vực này thường yêu cầu giữ chính xác thuật ngữ, con số, đơn vị đo và quan hệ giữa các sự kiện.",
    )
    anchor = add_picture_after(
        p91,
        image_template,
        FIGURE_DIR / "chapter1_taxonomy.png",
        6.2,
        "Phân loại bài toán tóm tắt văn bản",
        "Ba góc nhìn phân loại theo cách tạo đầu ra, số lượng tài liệu và mục tiêu sử dụng.",
    )
    add_caption_after(anchor, caption_template, "Hình 1.1: Ba góc nhìn phân loại bài toán tóm tắt văn bản")

    replace_text(p92, "1.4. Đặc thù tiếng Việt và ảnh hưởng đến bài toán tóm tắt")
    replace_text(
        p93,
        "Tiếng Việt đặt ra những yêu cầu riêng cho tiền xử lý, biểu diễn văn bản và đánh giá kết quả. Chương 1 chỉ xác định các vấn đề ở mức tổng quan; cơ chế kỹ thuật và ví dụ chi tiết được trình bày trong Chương 2.",
    )
    replace_text(
        p94,
        "Ranh giới từ không trùng hoàn toàn với khoảng trắng. Một từ có thể gồm nhiều tiếng, như “bệnh viện” hoặc “cấp cứu”. Nếu hệ thống coi từng tiếng là một từ độc lập, trọng số thống kê và biểu diễn ngữ nghĩa của câu có thể bị sai lệch.",
    )
    replace_text(
        p95,
        "Nhiều từ có nghĩa thay đổi theo ngữ cảnh. Từ “đường” có thể chỉ tuyến giao thông, chất tạo ngọt hoặc định hướng. Mô hình chỉ dựa trên đếm từ khó phân biệt các trường hợp này; biểu diễn theo ngữ cảnh phù hợp hơn khi cần so sánh ý nghĩa giữa các câu.",
    )
    replace_text(
        p96,
        "Tên riêng, chữ viết tắt, số liệu và đơn vị đo cần được giữ nhất quán. Một bản tóm tắt có thể thay đổi ý nghĩa nghiêm trọng nếu ghép sai tên chủ thể, bỏ dấu phủ định, nhầm mốc thời gian hoặc biến đổi con số.",
    )
    replace_text(
        p97,
        "Nguồn dữ liệu tiếng Việt có mức độ chuẩn hóa không đồng đều và khác biệt lớn giữa tin tức, mạng xã hội, văn bản hành chính và tài liệu chuyên ngành. Do đó, kết quả trên một bộ dữ liệu không tự động đại diện cho mọi lĩnh vực; việc đánh giá cần nêu rõ nguồn dữ liệu và phạm vi áp dụng.",
    )

    replace_text(p98, "1.5. Tổng quan nghiên cứu")
    replace_text(p99, "1.5.1. Tiến trình phát triển trên thế giới")
    replace_text(
        p100,
        "Giai đoạn đầu tập trung vào các dấu hiệu bề mặt. Luhn (1958) dùng tần suất từ để xác định đoạn có nhiều từ quan trọng [1]. Edmundson (1969) bổ sung vị trí câu, từ chỉ dẫn và từ xuất hiện trong tiêu đề [2]. Các phương pháp này dễ giải thích nhưng phụ thuộc mạnh vào đặc trưng thiết kế thủ công.",
    )
    replace_text(
        p101,
        "Năm 2004, Mihalcea và Tarau đề xuất TextRank, mô hình hóa văn bản dưới dạng đồ thị và xếp hạng các đơn vị ngôn ngữ bằng quan hệ tương đồng [3]. TextRank không cần dữ liệu gán nhãn và vẫn là một đường cơ sở phổ biến cho tóm tắt trích xuất.",
    )
    replace_text(
        p102,
        "Các mô hình sinh tóm tắt dựa trên mạng nơ-ron phát triển mạnh từ giữa thập niên 2010. Rush và cộng sự (2015) sử dụng Attention cho tóm tắt trừu tượng ở mức câu [4]. See và cộng sự (2017) kết hợp cơ chế sinh từ với sao chép từ nguồn, đồng thời dùng coverage để giảm lặp [5].",
    )
    replace_text(
        p103,
        "Transformer loại bỏ xử lý tuần tự bắt buộc và sử dụng Self-Attention để mô hình hóa quan hệ giữa các token [6]. Trên nền kiến trúc này, các mô hình tiền huấn luyện như BART [7] và T5 [8] tạo ra một khung chung cho nhiều tác vụ sinh văn bản, trong đó có tóm tắt.",
    )
    anchor = add_picture_after(
        p103,
        image_template,
        FIGURE_DIR / "chapter1_timeline.png",
        6.2,
        "Dòng thời gian nghiên cứu tóm tắt tự động",
        "Các mốc từ phương pháp thống kê, đồ thị, Attention, Transformer đến mô hình tiền huấn luyện cho tiếng Việt.",
    )
    add_caption_after(anchor, caption_template, "Hình 1.2: Các mốc tiêu biểu trong tiến trình phát triển phương pháp tóm tắt tự động")

    replace_text(p104, "1.5.2. Nghiên cứu và tài nguyên cho tiếng Việt")
    replace_text(
        p105,
        "Các bộ công cụ như VnCoreNLP [9] và Underthesea [10] cung cấp những chức năng nền tảng cho xử lý tiếng Việt, chẳng hạn tách từ, gán nhãn từ loại hoặc nhận diện thực thể. Chúng hỗ trợ chuẩn hóa đầu vào cho tóm tắt, nhưng chất lượng vẫn phụ thuộc loại văn bản và miền dữ liệu.",
    )
    replace_text(
        p106,
        "PhoBERT là mô hình ngôn ngữ tiền huấn luyện dành cho tiếng Việt, cung cấp biểu diễn theo ngữ cảnh và đạt kết quả mạnh trên nhiều tác vụ hiểu ngôn ngữ [11]. PhoBERT là mô hình kiểu encoder, vì vậy bản thân nó không trực tiếp sinh bản tóm tắt; có thể dùng vector của mô hình làm đầu vào cho xếp hạng hoặc phân cụm câu.",
    )
    replace_text(
        p107,
        "ViT5 xây dựng mô hình text-to-text cho tiếng Việt trên nền T5 và được đánh giá trên các tác vụ sinh ngôn ngữ, bao gồm tóm tắt trừu tượng [12]. Kiến trúc encoder-decoder cho phép tạo câu mới, nhưng chất lượng phụ thuộc dữ liệu huấn luyện, độ dài đầu vào và chiến lược giải mã.",
    )
    replace_text(p108, "1.5.3. So sánh các công trình tiêu biểu")
    replace_text(p109, "Bảng 1.1: So sánh tám công trình và mô hình tiêu biểu")
    replace_literature_table(document, literature_table, p109)

    replace_text(p110, "1.6. Khoảng trống nghiên cứu và định hướng của đề tài")
    replace_text(
        p111,
        "Các công trình được khảo sát thường khác nhau về dữ liệu, ngân sách độ dài và tiêu chí đánh giá. Vì vậy, vẫn thiếu một so sánh nhất quán giữa các hướng tiếp cận trên cùng đầu vào tiếng Việt, đồng thời xét chất lượng, chi phí tính toán và khả năng giải thích.",
    )
    replace_text(
        p112,
        "Câu hỏi nghiên cứu 1: Trên cùng dữ liệu tiếng Việt, chất lượng của TextRank, PhoBERT kết hợp phân cụm và ViT5 khác nhau như thế nào?",
    )
    replace_text(
        p113,
        "Câu hỏi nghiên cứu 2: Đánh đổi giữa chất lượng, tài nguyên tính toán, thời gian xử lý và khả năng truy vết của từng phương pháp là gì?",
    )
    replace_text(
        p114,
        "Câu hỏi nghiên cứu 3: Đặc thù tiếng Việt và loại văn bản ảnh hưởng thế nào đến tiền xử lý, biểu diễn câu và độ ổn định của kết quả?",
    )
    anchor = add_callout_after(
        p114,
        normal,
        "Định hướng của đề tài. ",
        "Đề tài dùng TextRank làm đường cơ sở, PhoBERT kết hợp K-Means/DBSCAN cho hướng trích xuất ngữ nghĩa và ViT5 cho hướng sinh tạo; sau đó so sánh trong cùng một khung đánh giá. Đây là định hướng nghiên cứu, không phải kết luận thực nghiệm.",
        fill="EAF3F8",
        accent="234E70",
    )
    anchor = add_text_after(anchor, heading3, "1.7. Tiểu kết Chương 1")
    add_text_after(
        anchor,
        normal,
        "Chương 1 đã xác định bài toán, yêu cầu chất lượng, các hướng phân loại, đặc thù tiếng Việt và tiến trình nghiên cứu. Từ khoảng trống được chỉ ra, đề tài lựa chọn TextRank, PhoBERT kết hợp phân cụm và ViT5 để khảo sát; cơ sở của các phương pháp được trình bày ở Chương 2.",
    )
    chapter2_heading.paragraph_format.page_break_before = True

    replace_text(
        p118,
        "Xử lý Ngôn ngữ Tự nhiên (Natural Language Processing - NLP) là một nhánh của Trí tuệ Nhân tạo, giao thoa giữa Khoa học Máy tính và Ngôn ngữ học tính toán. Có thể hình dung NLP là tập hợp các kỹ thuật giúp máy tính làm việc với lời nói và chữ viết: nhận biết từ, hiểu mối liên hệ giữa các câu, tìm thông tin quan trọng và tạo ra câu trả lời mới. Tóm tắt tự động là bài toán khó vì hệ thống phải đồng thời xác định ý chính, loại ý lặp, giữ đúng sự thật và tạo đầu ra đủ ngắn để người đọc tiết kiệm thời gian.",
    )
    anchor = p118
    anchor = add_picture_after(
        anchor,
        image_template,
        FIGURE_DIR / "chapter2_pipeline.png",
        6.2,
        "Quy trình tóm tắt văn bản",
        "Năm bước gồm văn bản thô, tiền xử lý, biểu diễn số, tóm tắt và đánh giá.",
    )
    add_caption_after(anchor, caption_template, "Hình 2.1: Quy trình tổng quát của một hệ thống tóm tắt văn bản tự động")

    replace_text(
        p123,
        "Rào cản về ranh giới từ (Word Boundary Ambiguity): Trong tiếng Việt, khoảng trắng chủ yếu ngăn cách các tiếng, không luôn ngăn cách các từ hoàn chỉnh. Vì vậy, cùng một chuỗi có thể được nhóm theo nhiều cách và tạo ra nghĩa khác nhau:",
    )
    replace_text(p124, "Cách hiểu A: Ông_già / đi / nhanh / quá. Nghĩa là một người lớn tuổi đang đi nhanh.")
    replace_text(p125, "Cách hiểu B: Ông / già_đi / nhanh / quá. Nghĩa là một người đang trở nên già nhanh.")
    replace_text(
        p126,
        "Con người dùng ngữ cảnh để chọn cách hiểu phù hợp. Máy tính cần một bước tách từ chuyên dụng; nếu nhóm sai từ ngay từ đầu thì TF-IDF, PhoBERT và các thuật toán tóm tắt phía sau đều nhận đầu vào sai.",
    )
    replace_text(p127, "2.1.3. Phân mảnh từ (Word Segmentation) và minh họa nhãn biên B/I")
    replace_text(
        p128,
        "Phân mảnh từ là quá trình xác định ranh giới giữa các từ trong một chuỗi âm tiết. Đầu vào là câu tiếng Việt chưa được phân tách theo đơn vị từ; đầu ra là dãy từ đơn và từ ghép đã được nhận diện. Do một từ tiếng Việt có thể gồm nhiều âm tiết cách nhau bằng khoảng trắng, việc xác định ranh giới phải dựa vào ngữ cảnh thay vì chỉ tách câu tại mỗi khoảng trắng.",
    )
    replace_text(
        p129,
        "Một cách biểu diễn bài toán tách từ là gán nhãn biên B/I. Với chuỗi âm tiết X = (x₁, x₂, ..., xₙ), mỗi vị trí nhận một nhãn yᵢ: B (Begin) báo hiệu bắt đầu một từ mới, còn I (Inside) báo hiệu âm tiết tiếp tục thuộc từ đang xét. Sau khi dự đoán nhãn cho toàn bộ chuỗi, các âm tiết mang nhãn I được ghép với âm tiết B đứng trước để tạo thành từ hoàn chỉnh.",
    )
    anchor = add_callout_after(
        p134,
        normal,
        "Ví dụ B/I. ",
        "Chuỗi “bệnh viện mở khu cấp cứu” có thể nhận nhãn bệnh(B) – viện(I) – mở(B) – khu(B) – cấp(B) – cứu(I). Sau khi ghép các vị trí B/I, đầu ra là: bệnh_viện / mở / khu / cấp_cứu.",
        fill="E9F6F2",
    )
    anchor = add_code_after(
        anchor,
        normal,
        "Mã minh họa công khai - tách từ tiếng Việt",
        "from underthesea import word_tokenize\n\n"
        "sentence = \"Bệnh viện mở khu cấp cứu mới\"\n"
        "tokens = word_tokenize(sentence, format=\"text\")\n"
        "print(tokens)",
        10,
    )
    put_picture(
        p135,
        FIGURE_DIR / "chapter2_preprocessing.png",
        6.2,
        "Tiền xử lý tiếng Việt",
        "Ví dụ văn bản thô được làm sạch, tách từ và gắn nhãn biên B/I.",
    )
    replace_text(p136, "Hình 2.2: Minh họa quá trình làm sạch và tách từ tiếng Việt")
    anchor = add_callout_after(
        p141,
        normal,
        "Ví dụ làm sạch. ",
        "Đầu vào: `<p> Bệnh viện A khai trương khu cấp cứu mới!!! Xem tại https://bv-a.vn </p>`. Sau bước làm sạch: `Bệnh viện A khai trương khu cấp cứu mới!`. Sau bước tách từ: `bệnh_viện A khai_trương khu cấp_cứu mới`.",
        fill="FFF7DD",
        accent="D69E2E",
    )
    replace_text(
        p143,
        "Stopwords (từ dừng) là những từ xuất hiện thường xuyên nhưng thường đóng góp ít vào việc phân biệt chủ đề, chẳng hạn “là”, “và”, “những”, “các”, “thì”, “mà”, “bị”, “được”. Không phải lúc nào từ dừng cũng vô nghĩa; việc loại bỏ chúng phù hợp với TextRank/TF-IDF, nhưng không nên áp dụng máy móc cho mô hình sinh văn bản vì có thể làm hỏng ngữ pháp.",
    )
    replace_text(
        p144,
        "Danh sách từ dừng cần được xây dựng phù hợp với ngôn ngữ và lĩnh vực của văn bản. Sau khi tách từ, các từ nằm trong danh sách này có thể được loại khỏi dữ liệu dùng để tính trọng số. Mức độ rút gọn phụ thuộc vào nội dung đầu vào, vì vậy không có một tỷ lệ loại bỏ cố định áp dụng cho mọi tài liệu.",
    )
    add_callout_after(
        p144,
        normal,
        "Ví dụ trước/sau. ",
        "“Bệnh viện đã và đang mở khu cấp cứu mới” → sau khi loại một số từ dừng còn các từ mang chủ đề rõ hơn như “bệnh_viện”, “mở”, “khu”, “cấp_cứu”, “mới”. Câu gốc vẫn được giữ riêng để xuất kết quả; danh sách rút gọn chỉ phục vụ tính toán.",
    )

    anchor = add_picture_after(
        p147,
        image_template,
        FIGURE_DIR / "chapter2_representation.png",
        6.2,
        "Các phương pháp biểu diễn văn bản",
        "So sánh Bag-of-Words, TF-IDF, Word2Vec và PhoBERT theo mức độ giữ ngữ cảnh.",
    )
    add_caption_after(anchor, caption_template, "Hình 2.3: Tiến trình phát triển từ đếm từ đến biểu diễn theo ngữ cảnh")
    anchor = add_callout_after(
        p158,
        normal,
        "Ví dụ tính TF-IDF đơn giản. ",
        "Giả sử có 3 câu. Từ `bệnh_viện` xuất hiện trong cả 3 câu nên IDF = log(3/3) = 0 và không giúp phân biệt câu nào. Từ `cấp_cứu` chỉ xuất hiện trong 1 câu nên IDF = log(3/1) ≈ 1,10; nếu từ này còn xuất hiện nhiều trong chính câu đó, trọng số TF-IDF sẽ cao. Nói ngắn gọn: TF-IDF ưu tiên từ vừa nổi bật trong một câu, vừa tương đối hiếm trong toàn bộ tập.",
        fill="FFF7DD",
        accent="D69E2E",
    )
    add_code_after(
        anchor,
        normal,
        "Mã minh họa công khai - tạo ma trận TF-IDF",
        "from sklearn.feature_extraction.text import TfidfVectorizer\n\n"
        "corpus = [\"bệnh_viện mở khu cấp_cứu\",\n"
        "          \"khu cấp_cứu hoạt_động liên_tục\"]\n"
        "X = TfidfVectorizer().fit_transform(corpus)\n"
        "print(X.shape)",
        13,
    )
    replace_text(
        p160,
        "Word2Vec được Mikolov và cộng sự giới thiệu năm 2013 [14]. Mô hình sử dụng mạng nơ-ron nông để học vector từ trong không gian liên tục, với hai kiến trúc chính là CBOW và Skip-gram. Số chiều của vector là một siêu tham số được lựa chọn khi huấn luyện, không phải một giá trị cố định cho mọi mô hình.",
    )
    replace_text(
        p165,
        "Hạn chế cốt lõi: Word2Vec là mô hình nhúng tĩnh, nên một từ chỉ có một vector dù ngữ cảnh thay đổi. Chẳng hạn từ “đường” trong “con đường đông xe”, “đường cát dùng để nấu ăn” và “đường lối phát triển” mang ba nghĩa khác nhau nhưng vẫn dùng chung một biểu diễn. Đây là lý do cần đến biểu diễn theo ngữ cảnh.",
    )
    replace_text(
        p167,
        "Sự ra đời của kiến trúc Transformer [6] đã thúc đẩy các mô hình biểu diễn theo ngữ cảnh. PhoBERT [11] là mô hình ngôn ngữ dành cho tiếng Việt, được phát triển dựa trên kiến trúc RoBERTa. Ở phiên bản Base, mỗi token được biểu diễn bằng vector 768 chiều; các vector token sau đó có thể được tổng hợp để tạo thành một vector đại diện cho toàn câu.",
    )
    replace_text(
        p168,
        "PhoBERT xem xét các từ ở cả bên trái và bên phải thông qua cơ chế Self-Attention. Thay vì đọc mỗi từ riêng lẻ, mô hình ước lượng từ nào trong câu cần được chú ý khi đang xử lý một token cụ thể. Công thức tổng quát là:",
    )
    replace_text(
        p170,
        "Trong đó Q (Query) có thể hiểu là “token hiện tại đang cần tìm thông tin gì”, K (Key) là “mỗi token khác có liên quan đến truy vấn đến đâu”, còn V (Value) là phần thông tin sẽ được truyền tiếp. Phép softmax biến các điểm liên quan thành trọng số; token liên quan hơn đóng góp nhiều hơn vào biểu diễn mới.",
    )
    replace_text(
        p171,
        "Nhờ cơ chế này, vector của từ “đường” trong ba câu ở mục trước sẽ khác nhau theo ngữ cảnh. Tuy nhiên, vector 768 chiều không có một trục đơn lẻ mang nghĩa “y tế” hay “giao thông”; ý nghĩa nằm trong toàn bộ cấu hình số và chỉ được diễn giải thông qua khoảng cách hoặc tác vụ phía sau.",
    )
    replace_text(
        p172,
        "Để biểu diễn cả câu bằng một vector duy nhất, có thể sử dụng Mean Pooling có mặt nạ: cộng vector của các token thực, bỏ qua token đệm (padding), rồi chia cho số token hợp lệ:",
    )
    add_callout_after(
        p173,
        normal,
        "Kết quả biểu diễn. ",
        "Nếu văn bản có n câu và mỗi câu được biểu diễn bằng một vector 768 chiều, kết quả thu được là ma trận có kích thước n × 768. Mỗi hàng tương ứng với một câu và có thể được dùng làm đầu vào cho các thuật toán phân cụm như K-Means hoặc DBSCAN.",
    )

    replace_text(
        p176,
        "TextRank là thuật toán xếp hạng dựa trên đồ thị và không cần dữ liệu gán nhãn [3]. Có thể hình dung mỗi câu là một người trong mạng lưới: một câu được coi là quan trọng khi nó giống nhiều câu khác, đặc biệt là giống các câu vốn cũng có điểm cao.",
    )
    add_text_after(
        p182,
        normal,
        "Ví dụ ở Hình 2.4 sử dụng năm câu về một khu cấp cứu. S1, S2 và S5 cùng nói về việc mở cơ sở, quy mô vận hành và tác động nên có nhiều cạnh mạnh. S4 chỉ nói về màu trang trí, ít liên quan đến ý chính nên có điểm thấp.",
    )
    put_picture(
        p183,
        FIGURE_DIR / "chapter2_textrank.png",
        6.2,
        "Đồ thị TextRank",
        "Năm câu được nối theo độ tương đồng, xếp hạng và chọn hai câu có điểm cao nhất.",
    )
    replace_text(p184, "Hình 2.4: Đồ thị tương đồng và quá trình chọn câu của TextRank")
    replace_text(
        p189,
        "d = 0,85 là hệ số giảm chấn (Damping Factor). Cách diễn giải trực quan: ở mỗi vòng lặp, phần lớn điểm số được truyền theo các cạnh tương đồng, còn một phần nhỏ giúp hệ thống không bị mắc kẹt trong một nhóm câu khép kín.",
    )
    replace_text(
        p191,
        "Thuật toán tiếp tục cập nhật điểm cho đến khi mức thay đổi giữa hai vòng lặp nhỏ hơn một ngưỡng ε. Trong thực tế, người ta cũng đặt một số vòng lặp tối đa để tránh trường hợp thuật toán mất quá nhiều thời gian khi chưa đạt điều kiện hội tụ.",
    )
    replace_text(p192, "|WS(Vᵢ)^(t+1) - WS(Vᵢ)^t| < ε")
    format_equation(p192)
    anchor = add_callout_after(
        p193,
        normal,
        "Từ điểm số đến bản tóm tắt. ",
        "Giả sử văn bản có 5 câu và cần lấy khoảng 30% nội dung, có thể chọn 2 câu đạt điểm cao nhất. Các câu được chọn nên được sắp xếp lại theo thứ tự xuất hiện trong văn bản gốc để bản tóm tắt không làm đảo lộn trình tự trình bày.",
        fill="E9F6F2",
    )
    add_code_after(
        anchor,
        normal,
        "Mã minh họa công khai - xếp hạng PageRank",
        "import networkx as nx\n\n"
        "G = nx.Graph()\n"
        "G.add_weighted_edges_from([(\"S1\", \"S2\", 0.8),\n"
        "                           (\"S2\", \"S3\", 0.6)])\n"
        "scores = nx.pagerank(G, alpha=0.85, weight=\"weight\")",
        15,
    )

    replace_text(
        p196,
        "TextRank đo tương đồng chủ yếu từ TF-IDF, nên hai câu dùng từ khác nhau nhưng cùng nghĩa có thể bị xem là ít liên quan. Một hướng khắc phục là sử dụng PhoBERT để biến mỗi câu thành vector có ngữ cảnh, sau đó phân cụm các vector nhằm nhận diện và bao quát nhiều nhóm ý trong văn bản.",
    )
    replace_text(
        p198,
        "K-Means chia toàn bộ câu vào K nhóm, trong đó K phải được xác định trước. Giá trị K có thể được lựa chọn dựa trên số chủ đề dự kiến hoặc các phương pháp đánh giá cụm như Elbow và Silhouette. Chẳng hạn, khi chọn K = 3 cho văn bản gồm 9 câu, thuật toán sẽ tìm ba nhóm câu gần nhau trong không gian vector.",
    )
    replace_text(
        p206,
        "Trong bài toán tóm tắt trích xuất, sau khi K-Means hội tụ, câu có vector gần tâm mỗi cụm nhất có thể được chọn làm câu đại diện. Cách làm này tăng cơ hội bao quát nhiều nhóm ý, nhưng không bảo đảm mỗi cụm luôn tương ứng với một chủ đề dễ gọi tên hoặc câu được chọn là câu tốt nhất về mặt diễn đạt.",
    )
    replace_text(
        p208,
        "DBSCAN (Density-Based Spatial Clustering of Applications with Noise) tạo cụm theo mật độ thay vì buộc mọi câu vào một trong K cụm. Thuật toán sử dụng hai tham số chính: bán kính lân cận ε và số điểm tối thiểu MinPts. Một vùng chỉ được mở rộng thành cụm khi có đủ số điểm nằm gần nhau theo hai điều kiện này.",
    )
    replace_text(
        p212,
        "DBSCAN có thể tạo bản tóm tắt rất ngắn khi chỉ hình thành ít cụm và loại nhiều câu nhiễu. Tuy nhiên, tỷ lệ nén không cố định: ε quá nhỏ có thể khiến nhiều câu bị xem là nhiễu, còn ε quá lớn có thể gom hầu hết câu vào cùng một cụm. Vì vậy, hai tham số ε và MinPts cần được lựa chọn phù hợp với đặc điểm dữ liệu.",
    )
    anchor = add_picture_after(
        p212,
        image_template,
        FIGURE_DIR / "chapter2_clustering.png",
        6.2,
        "So sánh K-Means và DBSCAN",
        "K-Means gán mọi câu vào một cụm, còn DBSCAN tạo cụm theo mật độ và có thể đánh dấu nhiễu.",
    )
    anchor = add_caption_after(anchor, caption_template, "Hình 2.5: Minh họa trực quan K-Means và DBSCAN trên vector câu")
    add_code_after(
        anchor,
        normal,
        "Mã minh họa công khai - phân cụm vector",
        "from sklearn.cluster import KMeans, DBSCAN\n\n"
        "kmeans_labels = KMeans(n_clusters=3, n_init=\"auto\").fit_predict(X)\n"
        "dbscan_labels = DBSCAN(eps=0.5, min_samples=2).fit_predict(X)",
        16,
    )

    replace_text(
        p216,
        "Các kiến trúc Seq2Seq dựa trên RNN thường đọc tuần tự và nén thông tin vào trạng thái ẩn [17]; cơ chế Attention giúp mô hình truy cập linh hoạt hơn tới các trạng thái của bộ mã hóa [18]. Transformer [6] thay thế xử lý tuần tự bắt buộc bằng Self-Attention, cho phép mỗi token tham chiếu trực tiếp tới các token liên quan và hỗ trợ tính toán song song tốt hơn.",
    )
    put_picture(
        p217,
        FIGURE_DIR / "chapter2_vit5.png",
        6.2,
        "Kiến trúc ViT5 và Beam Search",
        "Văn bản đi qua Encoder và Decoder; Beam Search giữ bốn hướng sinh tiềm năng trước khi chọn đầu ra.",
    )
    replace_text(p218, "Hình 2.6: Luồng Encoder–Decoder của ViT5 và minh họa Beam Search")
    replace_text(
        p219,
        "ViT5 xây dựng theo khung T5 (Text-to-Text Transfer Transformer) [12]. Encoder nhận văn bản nguồn và tạo biểu diễn theo ngữ cảnh; Decoder dùng biểu diễn đó để sinh từng token của bản tóm tắt. Vì đầu ra được viết lại, ViT5 có thể gộp thông tin nằm ở nhiều câu và tạo văn phong mạch lạc hơn, nhưng cũng có nguy cơ thêm chi tiết không có trong nguồn nếu mô hình hoặc dữ liệu huấn luyện chưa đủ tốt.",
    )
    add_callout_after(
        p219,
        normal,
        "Ví dụ trích xuất và sinh tạo. ",
        "Nguồn: “Khu cấp cứu có 30 giường. Khu này hoạt động 24/7. Bệnh viện bổ sung nhân sự.” TextRank/PhoBERT chỉ có thể lấy lại nguyên một hoặc nhiều câu nguồn. ViT5 có thể viết câu mới: “Bệnh viện đưa vào hoạt động khu cấp cứu 30 giường, phục vụ liên tục 24/7 và được bổ sung nhân sự.”",
        fill="E9F6F2",
    )
    replace_text(
        p221,
        "Khi Decoder sinh token mới, Greedy Search luôn chọn lựa chọn có xác suất cao nhất tại thời điểm hiện tại. Quyết định sớm này có thể dẫn đến một câu tổng thể kém hơn. Beam Search giữ đồng thời một số chuỗi ứng viên, mở rộng từng chuỗi ở bước tiếp theo và chỉ giữ lại các chuỗi có điểm tích lũy tốt nhất.",
    )
    replace_text(
        p222,
        "Số chuỗi ứng viên được giữ lại ở mỗi bước được gọi là độ rộng chùm (beam width). Chẳng hạn, với beam width bằng 4, thuật toán duy trì bốn hướng sinh có điểm cao nhất trước khi mở rộng chúng ở bước tiếp theo.",
    )
    replace_text(
        p226,
        "Một biện pháp hạn chế lặp là không cho phép n-gram đã xuất hiện được sinh lại nguyên vẹn. Ví dụ, với n = 3, một cụm ba token đã có trong chuỗi sẽ không được lặp lại. Ràng buộc này giảm hiện tượng lặp từ nhưng không thay thế việc kiểm tra tính đúng sự thật của bản tóm tắt.",
    )
    anchor = add_callout_after(
        p226,
        normal,
        "Cách đọc điểm Beam Search. ",
        "Mỗi nhánh có điểm là tổng log-xác suất các token đã sinh; Length Penalty điều chỉnh để chuỗi ngắn không tự động chiếm ưu thế chỉ vì có ít bước nhân xác suất hơn. Các xác suất trong Hình 2.6 chỉ là số giả định để giải thích cơ chế, không phải kết quả đo của mô hình.",
        fill="F3EEFA",
        accent="7663A7",
    )
    add_code_after(
        anchor,
        normal,
        "Mã minh họa công khai - kích hoạt Beam Search",
        "outputs = model.generate(\n"
        "    **inputs,\n"
        "    num_beams=4,\n"
        "    do_sample=False,\n"
        "    max_new_tokens=80,\n"
        ")",
        19,
    )

    replace_text(
        p229,
        "ROUGE (Recall-Oriented Understudy for Gisting Evaluation) là nhóm độ đo so sánh bản tóm tắt máy (Candidate C) với bản tóm tắt tham chiếu do con người viết (Reference R) [20]. Trước khi tính toán cho tiếng Việt, hai văn bản cần được chuẩn hóa và tách từ theo cùng một quy tắc để các đơn vị so sánh nhất quán.",
    )
    replace_text(
        p236,
        "ROUGE-L tìm chuỗi con chung dài nhất giữa Candidate và Reference mà vẫn giữ thứ tự token. Chỉ số này thưởng cho các cụm nội dung xuất hiện theo trật tự tương tự, nhưng không trực tiếp đo ngữ pháp, độ dễ đọc hay tính đúng sự thật.",
    )
    anchor = add_callout_after(
        p236,
        normal,
        "Ví dụ tính tay. ",
        "Reference: `bệnh_viện mở khu cấp_cứu mới` (5 token). Candidate: `bệnh_viện khai_trương khu cấp_cứu` (4 token). Có 3 unigram trùng nhau, nên Recall = 3/5 = 0,60; Precision = 3/4 = 0,75; F1 ≈ 0,67. Chỉ có một bigram trùng là `khu cấp_cứu`, nên ROUGE-2 thấp hơn.",
        fill="FFF7DD",
        accent="D69E2E",
    )
    anchor = add_picture_after(
        anchor,
        image_template,
        FIGURE_DIR / "chapter2_rouge.png",
        6.2,
        "Ví dụ ROUGE",
        "So sánh token giữa bản chuẩn và bản máy, kèm phép tính ROUGE-1, ROUGE-2 và ROUGE-L.",
    )
    anchor = add_caption_after(anchor, caption_template, "Hình 2.7: Ví dụ trực quan cách tính ROUGE trên một câu tiếng Việt")
    anchor = add_text_after(
        anchor,
        normal,
        "Giới hạn quan trọng: một bản tóm tắt có thể dùng từ đồng nghĩa và diễn đạt đúng nhưng ROUGE thấp; ngược lại, một câu sao chép nhiều từ có thể đạt ROUGE cao dù khó đọc hoặc sai quan hệ giữa các sự kiện. Vì vậy kết quả thực nghiệm cần kết hợp ROUGE với kiểm tra thủ công về độ mạch lạc và tính nhất quán sự thật.",
    )

    anchor = add_text_after(anchor, heading3, "2.7. Ví dụ tổng hợp: từ bài báo đến bản tóm tắt")
    anchor = add_text_after(
        anchor,
        normal,
        "Phần này nối các khái niệm của Chương 2 thành một ví dụ duy nhất. Các đầu ra dưới đây nhằm minh họa cơ chế và sự khác nhau giữa các phương pháp; chúng không thay thế số liệu thực nghiệm ở Chương 3.",
    )
    anchor = add_text_after(anchor, heading4, "2.7.1. Văn bản nguồn minh họa")
    anchor = add_callout_after(
        anchor,
        normal,
        "Bài báo 5 câu. ",
        "S1. Bệnh viện A khai trương khu cấp cứu mới.\nS2. Khu cấp cứu có 30 giường và hoạt động 24/7.\nS3. Bệnh viện bổ sung 12 bác sĩ và 20 điều dưỡng.\nS4. Khu vực chờ được trang trí màu xanh.\nS5. Cơ sở mới dự kiến giảm thời gian chờ của bệnh nhân.",
        fill="F4F6F8",
        accent="234E70",
    )
    anchor = add_text_after(anchor, heading4, "2.7.2. Cách bốn phương pháp xử lý cùng một đầu vào")
    anchor = add_callout_after(
        anchor,
        normal,
        "TextRank. ",
        "Tạo vector TF-IDF, nối các câu tương đồng, xếp hạng và lấy khoảng 30% số câu. Đầu ra là các câu gốc có điểm cao, chẳng hạn S1 và S2; câu chữ không được viết lại.",
        fill="EAF3F8",
        accent="4C78A8",
    )
    anchor = add_callout_after(
        anchor,
        normal,
        "PhoBERT + K-Means. ",
        "Giả sử chọn K = 3, các câu được gom thành ba nhóm theo mức độ gần nhau về ngữ nghĩa. Câu gần tâm của mỗi cụm được chọn làm đại diện, nên đầu ra có thể bao quát các nhóm ý như mở cơ sở, bổ sung nhân sự và tác động dự kiến.",
        fill="FFF7DD",
        accent="D69E2E",
    )
    anchor = add_callout_after(
        anchor,
        normal,
        "PhoBERT + DBSCAN. ",
        "Các câu tạo thành vùng có mật độ cao được gom cụm; câu đứng riêng có thể bị coi là nhiễu. Nếu S4 ít liên quan tới các câu còn lại, nó có thể bị loại. Số câu đầu ra phụ thuộc vào ε và MinPts nên không cố định từ trước.",
        fill="E9F6F2",
        accent="1F8A70",
    )
    anchor = add_callout_after(
        anchor,
        normal,
        "ViT5. ",
        "Encoder đọc toàn bộ năm câu, Decoder có thể tạo câu mới: “Bệnh viện A đưa vào hoạt động khu cấp cứu 30 giường, phục vụ 24/7 và bổ sung 32 nhân viên y tế nhằm giảm thời gian chờ.” Đầu ra gọn và liền mạch, nhưng cần đối chiếu lại các con số với nguồn.",
        fill="F3EEFA",
        accent="7663A7",
    )
    anchor = add_text_after(anchor, heading4, "2.7.3. Chọn phương pháp theo nhu cầu")
    anchor = add_bullet_after(anchor, bullet, "Chọn TextRank khi cần tốc độ cao, chạy tốt trên CPU và ưu tiên giữ nguyên câu nguồn.")
    anchor = add_bullet_after(anchor, bullet, "Chọn PhoBERT + K-Means khi cần bao quát nhiều nhóm ý và chấp nhận chi phí tạo embedding cao hơn.")
    anchor = add_bullet_after(anchor, bullet, "Chọn PhoBERT + DBSCAN khi muốn phát hiện nhóm ý theo mật độ và loại câu lạc đề, đồng thời có dữ liệu để hiệu chỉnh tham số.")
    anchor = add_bullet_after(anchor, bullet, "Chọn ViT5 khi ưu tiên bản tóm tắt tự nhiên, có hạ tầng mô hình sinh và có bước kiểm tra sai lệch thông tin.")
    anchor = add_text_after(anchor, heading3, "2.8. Tiểu kết Chương 2")
    add_text_after(
        anchor,
        normal,
        "Chương 2 đã trình bày cơ sở lý thuyết của tiền xử lý, biểu diễn văn bản, TextRank, phân cụm PhoBERT, ViT5 và ROUGE. Mỗi phương pháp có ưu điểm, hạn chế và phạm vi sử dụng riêng.",
    )

    for separator in separators:
        clear_separator(separator)

    page_break_before_chapter3._element.getparent().remove(page_break_before_chapter3._element)
    chapter3_heading.paragraph_format.page_break_before = True
    post_results_blank._element.getparent().remove(post_results_blank._element)
    page_break_before_conclusion._element.getparent().remove(page_break_before_conclusion._element)
    conclusion_heading.paragraph_format.page_break_before = True

    equation_table = document.tables[1]
    equation_table._element.getparent().remove(equation_table._element)

    core = document.core_properties
    core.title = "Báo cáo tổng kết - Chương 1 viết lại và Chương 2 bổ sung minh họa"
    core.subject = "Tóm tắt văn bản tiếng Việt"
    core.comments = "Chương 1 được viết lại theo hướng học thuật; Chương 2 được bổ sung hình minh họa; trích dẫn và tài liệu tham khảo được chuẩn hóa theo IEEE."

    ieee_references = (
        "[1] H. P. Luhn, “The automatic creation of literature abstracts,” IBM J. Res. Develop., vol. 2, no. 2, pp. 159–165, Apr. 1958.",
        "[2] H. P. Edmundson, “New methods in automatic extracting,” J. ACM, vol. 16, no. 2, pp. 264–285, Apr. 1969, doi: 10.1145/321510.321519.",
        "[3] R. Mihalcea and P. Tarau, “TextRank: Bringing order into texts,” in Proc. 2004 Conf. Empirical Methods Natural Language Processing (EMNLP), Barcelona, Spain, 2004, pp. 404–411.",
        "[4] A. M. Rush, S. Chopra, and J. Weston, “A neural attention model for abstractive sentence summarization,” in Proc. 2015 Conf. Empirical Methods Natural Language Processing (EMNLP), Lisbon, Portugal, 2015, pp. 379–389, doi: 10.18653/v1/D15-1044.",
        "[5] A. See, P. J. Liu, and C. D. Manning, “Get to the point: Summarization with pointer-generator networks,” in Proc. 55th Annu. Meeting Assoc. Comput. Linguistics (ACL), Vancouver, BC, Canada, 2017, pp. 1073–1083, doi: 10.18653/v1/P17-1099.",
        "[6] A. Vaswani et al., “Attention is all you need,” in Adv. Neural Inf. Process. Syst. 30 (NeurIPS), Long Beach, CA, USA, 2017, pp. 5998–6008.",
        "[7] M. Lewis et al., “BART: Denoising sequence-to-sequence pre-training for natural language generation, translation, and comprehension,” in Proc. 58th Annu. Meeting Assoc. Comput. Linguistics (ACL), Online, 2020, pp. 7871–7880, doi: 10.18653/v1/2020.acl-main.703.",
        "[8] C. Raffel et al., “Exploring the limits of transfer learning with a unified text-to-text transformer,” J. Mach. Learn. Res., vol. 21, no. 140, pp. 1–67, 2020.",
        "[9] T. Vu, D. Q. Nguyen, D. Q. Nguyen, M. Dras, and M. Johnson, “VnCoreNLP: A Vietnamese natural language processing toolkit,” in Proc. NAACL-HLT Demonstrations, New Orleans, LA, USA, 2018, pp. 56–60, doi: 10.18653/v1/N18-5012.",
        "[10] Underthesea Development Team, “Underthesea documentation: Word tokenization.” Accessed: Aug. 25, 2026. [Online]. Available: https://underthesea.readthedocs.io/en/latest/readme.html",
        "[11] D. Q. Nguyen and A. T. Nguyen, “PhoBERT: Pre-trained language models for Vietnamese,” in Findings Assoc. Comput. Linguistics: EMNLP 2020, Online, 2020, pp. 1037–1042, doi: 10.18653/v1/2020.findings-emnlp.92.",
        "[12] L. Phan, H. Tran, H. Nguyen, and T. H. Trinh, “ViT5: Pretrained text-to-text transformer for Vietnamese language generation,” in Proc. NAACL-HLT Student Research Workshop, Seattle, WA, USA, 2022, pp. 136–142, doi: 10.18653/v1/2022.naacl-srw.18.",
        "[13] Scikit-learn Developers, “TfidfVectorizer,” scikit-learn documentation. Accessed: Aug. 25, 2026. [Online]. Available: https://scikit-learn.org/stable/modules/generated/sklearn.feature_extraction.text.TfidfVectorizer.html",
        "[14] T. Mikolov, K. Chen, G. Corrado, and J. Dean, “Efficient estimation of word representations in vector space,” arXiv:1301.3781, 2013, doi: 10.48550/arXiv.1301.3781.",
        "[15] NetworkX Developers, “pagerank,” NetworkX documentation. Accessed: Aug. 25, 2026. [Online]. Available: https://networkx.org/documentation/stable/reference/algorithms/generated/networkx.algorithms.link_analysis.pagerank_alg.pagerank.html",
        "[16] Scikit-learn Developers, “Clustering,” scikit-learn user guide. Accessed: Aug. 25, 2026. [Online]. Available: https://scikit-learn.org/stable/modules/clustering.html",
        "[17] K. Cho et al., “Learning phrase representations using RNN encoder-decoder for statistical machine translation,” in Proc. 2014 Conf. Empirical Methods Natural Language Processing (EMNLP), Doha, Qatar, 2014, pp. 1724–1734, doi: 10.3115/v1/D14-1179.",
        "[18] D. Bahdanau, K. Cho, and Y. Bengio, “Neural machine translation by jointly learning to align and translate,” in Proc. 3rd Int. Conf. Learn. Represent. (ICLR), San Diego, CA, USA, 2015.",
        "[19] Hugging Face, “Generation,” Transformers documentation. Accessed: Aug. 25, 2026. [Online]. Available: https://huggingface.co/docs/transformers/main_classes/text_generation",
        "[20] C.-Y. Lin, “ROUGE: A package for automatic evaluation of summaries,” in Text Summarization Branches Out: Proc. ACL-04 Workshop, Barcelona, Spain, 2004, pp. 74–81.",
    )
    for paragraph, reference in zip(reference_paragraphs, ieee_references[: len(reference_paragraphs)]):
        replace_text(paragraph, reference)
        format_reference_entry(paragraph)
    for reference in ieee_references[len(reference_paragraphs) :]:
        reference_anchor = add_text_after(reference_anchor, normal, reference)
        format_reference_entry(reference_anchor)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    enhance(args.input, args.output)
    print(args.output)


if __name__ == "__main__":
    main()
